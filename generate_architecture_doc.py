"""
生成服务器架构详细构建说明Word文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import datetime

def create_architecture_document():
    """创建架构文档"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # ==================== 封面页 ====================
    title = doc.add_heading('服务器架构详细构建说明', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('阿里云ECS服务器 (47.103.108.47)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')
    
    # 基本信息表
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ['项目名称', '软件设计项目 - AI宠物美妆镜'],
        ['服务器地址', '47.103.108.47 (阿里云ECS)'],
        ['操作系统', 'Alibaba Cloud Linux 3.2104 U12'],
        ['文档版本', 'v1.0'],
        ['生成日期', datetime.datetime.now().strftime('%Y年%m月%d日')]
    ]
    
    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ==================== 目录 ====================
    doc.add_heading('目 录', level=1)
    
    toc_items = [
        '一、基础环境层',
        '二、中间件层', 
        '三、应用运行时层',
        '四、CI/CD自动化层',
        '五、应用部署层',
        '六、安全防护层',
        '七、监控与日志层',
        '八、完整部署流程',
        '九、关键技术栈总结',
        '十、注意事项与建议'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ==================== 第一章 ====================
    doc.add_heading('一、基础环境层', level=1)
    
    doc.add_heading('1.1 操作系统信息', level=2)
    os_table = doc.add_table(rows=6, cols=2)
    os_table.style = 'Table Grid'
    
    os_info = [
        ['配置项', '详细信息'],
        ['操作系统', 'Alibaba Cloud Linux 3.2104 U12 (OpenAnolis Edition)'],
        ['内核版本', '5.10.134-19.1.al8.x86_64'],
        ['系统架构', 'x86_64'],
        ['安装方式', '阿里云ECS实例（预装系统镜像）'],
        ['主机名', 'iZuf66w1hgu6setio3rcabZ']
    ]
    
    for i, row_data in enumerate(os_info):
        row = os_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading('1.2 核心服务管理', level=2)
    doc.add_paragraph('使用Systemd进行服务管理，以下服务设置为开机自启：')
    
    services = [
        ('docker.service', 'Docker容器引擎'),
        ('redis.service', 'Redis缓存数据库'),
        ('nginx.service', 'Nginx反向代理（宝塔管理）'),
        ('gunicorn-flask-data-server.service', 'Flask应用服务'),
        ('jenkins.service', 'CI/CD自动化部署'),
        ('firewalld.service', '防火墙服务')
    ]
    
    svc_table = doc.add_table(rows=len(services)+1, cols=2)
    svc_table.style = 'Light Shading Accent 1'
    header_row = svc_table.rows[0]
    header_row.cells[0].text = '服务名称'
    header_row.cells[1].text = '功能描述'
    header_row.cells[0].paragraphs[0].runs[0].font.bold = True
    header_row.cells[1].paragraphs[0].runs[0].font.bold = True
    
    for i, (svc_name, desc) in enumerate(services, 1):
        row = svc_table.rows[i]
        row.cells[0].text = svc_name
        row.cells[1].text = desc
    
    doc.add_page_break()
    
    # ==================== 第二章 ====================
    doc.add_heading('二、中间件层', level=1)
    
    doc.add_heading('2.1 Docker容器引擎', level=2)
    doc.add_paragraph('版本：Docker 26.1.3 + Containerd')
    
    doc.add_paragraph('\n运行的容器：')
    containers = [
        ('mirror_sim_1', 'AI宠物美妆镜传感器模拟器 #1', '运行中'),
        ('mirror_sim_2', 'AI宠物美妆镜传感器模拟器 #2', '运行中'),
        ('mirror_sim_3', 'AI宠物美妆镜传感器模拟器 #3', '运行中'),
        ('redis-server', 'Redis主服务', '已停止'),
        ('redis-queue', 'Redis队列服务', '已创建未启动')
    ]
    
    container_table = doc.add_table(rows=len(containers)+1, cols=3)
    container_table.style = 'Medium Grid 1 Accent 1'
    header = container_table.rows[0]
    header.cells[0].text = '容器名称'
    header.cells[1].text = '功能说明'
    header.cells[2].text = '状态'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, (name, func, status) in enumerate(containers, 1):
        row = container_table.rows[i]
        row.cells[0].text = name
        row.cells[1].text = func
        row.cells[2].text = status
    
    doc.add_heading('2.2 Redis缓存数据库', level=2)
    doc.add_paragraph('版本：Redis 6.2.20')
    
    redis_config = [
        ['配置项', '配置值', '说明'],
        ['bind', '0.0.0.0', '允许所有IP访问'],
        ['port', '6379', '默认端口'],
        ['protected-mode', 'yes', '启用保护模式'],
        ['daemonize', 'no', '由systemd管理进程'],
        ['dir', '/var/lib/redis', '数据持久化目录'],
        ['databases', '16', '数据库数量']
    ]
    
    redis_table = doc.add_table(rows=len(redis_config), cols=3)
    redis_table.style = 'Table Grid'
    for i, row_data in enumerate(redis_config):
        row = redis_table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_heading('2.3 Nginx反向代理', level=2)
    doc.add_paragraph('版本：Nginx 1.28.1（通过宝塔面板管理）')
    doc.add_paragraph('安装路径：/www/server/nginx/')
    
    nginx_features = [
        '自动检测CPU核心数配置worker进程',
        '最大连接数：51200',
        '启用GZIP压缩提升传输效率',
        '最大上传文件大小：50MB',
        '监听80端口（HTTP）',
        '监听8888端口（phpMyAdmin管理界面，仅本地访问）'
    ]
    
    for feature in nginx_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 第三章 ====================
    doc.add_heading('三、应用运行时层', level=1)
    
    doc.add_heading('3.1 Python环境', level=2)
    doc.add_paragraph('版本：Python 3.6.8（系统自带）')
    
    doc.add_paragraph('\n关键依赖包：')
    packages = [
        ('Flask', '2.0.3', 'Web框架'),
        ('gunicorn', '21.2.0', 'WSGI服务器'),
        ('gevent', '22.10.2', '异步协程库'),
        ('redis', '4.3.6', 'Redis客户端'),
        ('PyJWT', '1.6.1', 'JWT认证'),
        ('cryptography', '3.2.1', '加密库'),
        ('Flask-Cors', '5.0.0', '跨域支持'),
        ('Flask-Limiter', '1.5', '速率限制')
    ]
    
    pkg_table = doc.add_table(rows=len(packages)+1, cols=3)
    pkg_table.style = 'Light List Accent 1'
    header = pkg_table.rows[0]
    header.cells[0].text = '包名'
    header.cells[1].text = '版本'
    header.cells[2].text = '用途'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, (pkg, ver, usage) in enumerate(packages, 1):
        row = pkg_table.rows[i]
        row.cells[0].text = pkg
        row.cells[1].text = ver
        row.cells[2].text = usage
    
    doc.add_heading('3.2 Gunicorn应用服务器', level=2)
    doc.add_paragraph('配置文件：/root/course-project/week5/data-server/data-server/config/gunicorn_config.py')
    
    gunicorn_config = [
        ['配置项', '配置值', '说明'],
        ['bind', '0.0.0.0:5000', '监听地址和端口'],
        ['workers', 'CPU*2+1', 'Worker进程数量'],
        ['worker_class', 'gevent', '异步Worker类型'],
        ['threads', '2', '每个Worker的线程数'],
        ['max_requests', '1000', '最大请求数后重启'],
        ['timeout', '120', '超时时间（秒）'],
        ['preload_app', 'True', '预加载应用节省内存'],
        ['backlog', '2048', 'TCP队列长度']
    ]
    
    gun_table = doc.add_table(rows=len(gunicorn_config), cols=3)
    gun_table.style = 'Table Grid'
    for i, row_data in enumerate(gunicorn_config):
        row = gun_table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ==================== 第四章 ====================
    doc.add_heading('四、CI/CD自动化层', level=1)
    
    doc.add_heading('4.1 Jenkins持续集成', level=2)
    doc.add_paragraph('工作空间：/var/lib/jenkins/')
    
    jenkins_config = [
        ['配置项', '配置值'],
        ['任务名称', 'software-design-project'],
        ['Git仓库', 'https://github.com/MOONFISH2233/software-design-project.git'],
        ['分支', 'master'],
        ['触发器', 'GitHub Push事件自动触发'],
        ['Pipeline文件', 'Jenkinsfile']
    ]
    
    jenkins_table = doc.add_table(rows=len(jenkins_config), cols=2)
    jenkins_table.style = 'Medium Shading 1 Accent 1'
    for i, row_data in enumerate(jenkins_config):
        row = jenkins_table.rows[i]
        row.cells[0].text = row_data[0]
        row.cells[1].text = row_data[1]
        if i == 0:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph('\nJenkins Pipeline流程：')
    pipeline_stages = [
        '拉取代码：从GitHub克隆最新代码',
        '代码检查：执行Python语法检查（py_compile）',
        '安装依赖：pip3 install -r requirements.txt',
        '一键部署：执行deploy.sh部署脚本',
        '健康检查：curl测试API接口可用性'
    ]
    
    for stage in pipeline_stages:
        doc.add_paragraph(stage, style='List Number')
    
    doc.add_page_break()
    
    # ==================== 第五章 ====================
    doc.add_heading('五、应用部署层', level=1)
    
    doc.add_heading('5.1 项目结构', level=2)
    doc.add_paragraph('主部署路径：/root/course-project/week5/data-server/data-server/')
    
    project_files = [
        'app.py - Flask主应用（30KB+）',
        'config/gunicorn_config.py - Gunicorn配置',
        'logs/ - 日志目录（访问日志、错误日志、应用日志）',
        'data/ - 数据存储目录',
        'security/ - 安全配置（API Keys、用户信息）',
        'static/ - 静态文件（测试仪表板）',
        'tests/ - 测试脚本（自动化测试、压力测试）',
        'scripts/ - 部署脚本（deploy.sh、monitor_performance.sh）',
        'docs/ - 项目文档（API文档、部署指南）'
    ]
    
    for file_item in project_files:
        doc.add_paragraph(file_item, style='List Bullet')
    
    doc.add_heading('5.2 Flask应用架构', level=2)
    
    api_endpoints = [
        ('POST /api/data/receive', '接收传感器数据'),
        ('GET /api/data/query', '查询历史数据'),
        ('POST /api/auth/login', '用户登录获取Token'),
        ('GET /api/health', '健康检查'),
        ('POST /api/receive/secure', '加密数据接收'),
        ('POST /api/encrypt', '数据加密'),
        ('POST /api/decrypt', '数据解密')
    ]
    
    doc.add_paragraph('\n主要API端点：')
    api_table = doc.add_table(rows=len(api_endpoints)+1, cols=2)
    api_table.style = 'Light Grid Accent 1'
    header = api_table.rows[0]
    header.cells[0].text = '接口路径'
    header.cells[1].text = '功能说明'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, (path, desc) in enumerate(api_endpoints, 1):
        row = api_table.rows[i]
        row.cells[0].text = path
        row.cells[1].text = desc
    
    doc.add_page_break()
    
    # ==================== 第六章 ====================
    doc.add_heading('六、安全防护层', level=1)
    
    doc.add_heading('6.1 防火墙配置', level=2)
    doc.add_paragraph('工具：firewalld')
    
    ports = [
        ('22/tcp', 'SSH远程访问'),
        ('80/tcp', 'HTTP服务'),
        ('443/tcp', 'HTTPS服务'),
        ('5000/tcp', 'Flask应用'),
        ('6379/tcp', 'Redis数据库'),
        ('8080/tcp', '备用端口'),
        ('8888/tcp', '宝塔面板')
    ]
    
    doc.add_paragraph('\n开放端口：')
    port_table = doc.add_table(rows=len(ports)+1, cols=2)
    port_table.style = 'Table Grid'
    header = port_table.rows[0]
    header.cells[0].text = '端口'
    header.cells[1].text = '用途'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, (port, usage) in enumerate(ports, 1):
        row = port_table.rows[i]
        row.cells[0].text = port
        row.cells[1].text = usage
    
    doc.add_heading('6.2 应用安全', level=2)
    
    security_features = [
        'JWT认证：Token有效期24小时',
        '速率限制：100次/分钟，1000次/小时',
        'AES加密：敏感数据加密传输',
        'CORS控制：限制跨域访问来源',
        'API Key认证：额外的身份验证层',
        '日志脱敏：不记录敏感信息'
    ]
    
    for feature in security_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 第七章 ====================
    doc.add_heading('七、监控与日志层', level=1)
    
    doc.add_heading('7.1 日志系统', level=2)
    
    log_locations = [
        ('Gunicorn访问日志', '/root/course-project/week5/data-server/logs/gunicorn_access.log'),
        ('Gunicorn错误日志', '/root/course-project/week5/data-server/logs/gunicorn_error.log'),
        ('应用日志', '/root/course-project/week5/data-server/logs/server_YYYYMMDD.log'),
        ('错误日志', '/root/course-project/week5/data-server/logs/error_YYYYMMDD.log'),
        ('Nginx日志', '/www/wwwlogs/nginx_error.log')
    ]
    
    doc.add_paragraph('\n日志位置：')
    log_table = doc.add_table(rows=len(log_locations)+1, cols=2)
    log_table.style = 'Light Shading Accent 1'
    header = log_table.rows[0]
    header.cells[0].text = '日志类型'
    header.cells[1].text = '文件路径'
    for cell in header.cells:
        cell.paragraphs[0].runs[0].font.bold = True
    
    for i, (log_type, path) in enumerate(log_locations, 1):
        row = log_table.rows[i]
        row.cells[0].text = log_type
        row.cells[1].text = path
    
    doc.add_paragraph('\n日志特性：')
    log_features = [
        '异步处理（独立线程，不阻塞主流程）',
        'JSON格式（便于ELK等日志分析系统处理）',
        '日志轮转（50MB切割，保留10个备份）',
        '采样过滤（INFO级别10%采样，降低存储压力）',
        '分级存储（DEBUG/INFO/WARNING/ERROR分开记录）'
    ]
    
    for feature in log_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_page_break()
    
    # ==================== 第八章 ====================
    doc.add_heading('八、完整部署流程', level=1)
    
    doc.add_heading('8.1 手动部署步骤', level=2)
    
    deploy_steps = [
        '克隆代码：cd /root/course-project && git clone https://github.com/MOONFISH2233/software-design-project.git',
        '安装依赖：cd week5/data-server/data-server && pip3 install -r requirements.txt',
        '配置Gunicorn：cp config/gunicorn_config.py.example config/gunicorn_config.py',
        '创建目录：mkdir -p logs data security',
        '配置Systemd：复制service文件到/etc/systemd/system/',
        '重载配置：systemctl daemon-reload',
        '启动服务：systemctl start gunicorn-flask-data-server',
        '设置自启：systemctl enable gunicorn-flask-data-server',
        '启动Docker：cd /root && docker-compose up -d',
        '验证部署：curl http://localhost:5000/api/health'
    ]
    
    for i, step in enumerate(deploy_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('8.2 自动化部署（Jenkins）', level=2)
    
    auto_deploy_flow = [
        '推送代码到GitHub master分支',
        'GitHub Webhook自动触发Jenkins任务',
        'Jenkins拉取最新代码',
        '执行Python语法检查',
        '安装或更新依赖包',
        '执行deploy.sh部署脚本',
        '健康检查验证服务可用性',
        '发送部署结果通知'
    ]
    
    for step in auto_deploy_flow:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # ==================== 第九章 ====================
    doc.add_heading('九、关键技术栈总结', level=1)
    
    tech_stack = [
        ['层级', '技术', '版本', '用途'],
        ['操作系统', 'Alibaba Cloud Linux', '3.2104', '基础运行环境'],
        ['容器化', 'Docker', '26.1.3', '应用隔离与部署'],
        ['缓存', 'Redis', '6.2.20', '缓存/消息队列'],
        ['Web服务器', 'Nginx', '1.28.1', '反向代理/静态资源'],
        ['应用服务器', 'Gunicorn', '21.2.0', 'Python WSGI服务器'],
        ['Web框架', 'Flask', '2.0.3', 'RESTful API开发'],
        ['异步库', 'gevent', '22.10.2', '协程并发'],
        ['CI/CD', 'Jenkins', 'Latest', '自动化部署'],
        ['认证', 'PyJWT', '1.6.1', 'Token认证'],
        ['加密', 'cryptography', '3.2.1', 'AES加密'],
        ['限流', 'Flask-Limiter', '1.5', 'API速率限制']
    ]
    
    tech_table = doc.add_table(rows=len(tech_stack), cols=4)
    tech_table.style = 'Medium Grid 3 Accent 1'
    for i, row_data in enumerate(tech_stack):
        row = tech_table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].font.bold = True
                row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== 第十章 ====================
    doc.add_heading('十、注意事项与建议', level=1)
    
    warnings = [
        '密码安全：服务器密码为 @Dierzu999，建议定期更换',
        'Redis安全：当前绑定 0.0.0.0，建议配置密码认证增强安全性',
        'Python版本：3.6.8较老，建议升级到3.9+以获得更好的性能和安全性',
        'MySQL状态：MySQL服务未运行，如需使用需手动启动',
        'RabbitMQ：未安装，当前使用Redis作为消息队列替代方案',
        '备份策略：建议定期备份 /root/course-project/ 和 /var/lib/redis/ 目录',
        '监控告警：建议配置Prometheus+Grafana实现可视化监控和告警',
        '日志归档：建议配置日志归档策略，避免磁盘空间耗尽',
        'SSL证书：如需要HTTPS，可申请Let\'s Encrypt免费证书',
        '性能优化：可考虑使用Redis缓存热点数据，减少数据库查询'
    ]
    
    for warning in warnings:
        p = doc.add_paragraph(warning, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.3)
    
    # 添加页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'服务器架构详细构建说明 | 生成时间：{datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    filename = f'服务器架构详细构建说明_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(filename)
    print(f'✅ Word文档已生成：{filename}')
    return filename

if __name__ == '__main__':
    try:
        filename = create_architecture_document()
        print(f'\n📄 文档位置：d:\\学习\\软件设计\\{filename}')
        print(f'📊 文档包含：10个章节，完整的架构说明')
        print(f'✨ 格式化：表格、列表、标题层次清晰')
    except Exception as e:
        print(f'❌ 生成失败：{str(e)}')
        import traceback
        traceback.print_exc()
