from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

def create_feishu_report():
    """创建飞书汇报文档"""
    
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('第六周任务完成情况汇报 - MongoDB数据库集成', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 版本信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('团队名称：').bold = True
    p.add_run('软件开发小组\n')
    p.add_run('汇报日期：').bold = True
    p.add_run(f'{datetime.now().strftime("%Y-%m-%d")}\n')
    p.add_run('Git分支：').bold = True
    p.add_run('week6')
    
    doc.add_page_break()
    
    # ==================== 一、本周工作概述 ====================
    doc.add_heading('一、本周工作概述', level=1)
    
    overview = '''
本周（第六周）完成了数据库存储方案的升级，将原有的文件存储改造为MongoDB数据库存储，并建立了完整的运维体系。

核心成果：
✅ 安装并配置 MongoDB 6.0.27
✅ 完成完整的数据库设计（3个集合 + 11个索引）
✅ 改造 module_writer.py 支持 MongoDB 双写模式
✅ 建立自动化备份机制（每日凌晨2点）
✅ 部署实时监控告警系统（每5分钟）
✅ 完成性能对比测试（写入QPS提升34.9%）

交付物：
• 4份专业文档（Word/Excel格式）
• 4个运维工具脚本
• 完整的Git代码仓库（week6分支）
'''
    doc.add_paragraph(overview)
    
    doc.add_page_break()
    
    # ==================== 二、详细完成情况 ====================
    doc.add_heading('二、详细完成情况', level=1)
    
    # 2.1 MongoDB安装与配置
    doc.add_heading('2.1 MongoDB安装与配置', level=2)
    
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = 'Medium Grid 3 Accent 1'
    hdr_cells = table1.rows[0].cells
    headers = ['项目', '详情']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    mongo_info = [
        ['版本号', 'MongoDB 6.0.27'],
        ['服务状态', 'Active (running) ✅'],
        ['监听端口', '27017'],
        ['数据目录', '/var/lib/mongo'],
        ['日志目录', '/var/log/mongodb'],
        ['配置文件', '/etc/mongod.conf']
    ]
    
    for item in mongo_info:
        row_cells = table1.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
    
    doc.add_paragraph('\n验证命令：')
    doc.add_paragraph('systemctl status mongod', style='No Spacing')
    
    # 2.2 数据库设计
    doc.add_heading('2.2 数据库设计', level=2)
    
    db_design = '''
设计了3个集合用于存储不同类型的传感器数据：

1️⃣ skin_sensor（皮肤传感器）
   • 字段：device_id, moisture, oiliness, temperature, timestamp等
   • 索引：device_id+timestamp复合索引、received_at时间索引
   
2️⃣ environment_sensor（环境传感器）
   • 字段：device_id, pm25, co2, humidity, location等
   • 索引：location地理空间索引、timestamp时间索引
   
3️⃣ device_status（设备状态）
   • 字段：device_id, status, battery_level, signal_strength等
   • 索引：device_id+last_heartbeat复合索引

总计创建11个索引优化查询性能。
'''
    doc.add_paragraph(db_design)
    
    # 2.3 代码改造
    doc.add_heading('2.3 代码改造', level=2)
    
    code_changes = '''
改造了 module_writer.py，实现灵活的存储架构：

🔧 三种存储模式：
• STORAGE_MODE = 'mongodb'  → 纯MongoDB模式（推荐生产环境）
• STORAGE_MODE = 'file'     → 纯文件模式（降级方案）
• STORAGE_MODE = 'both'     → 双写模式（过渡期使用）

✨ 技术亮点：
• 自动故障转移：MongoDB失败时自动降级到文件存储
• 连接池管理：maxPoolSize=50, minPoolSize=10
• 批量写入优化：使用bulk_write提高性能
• 向后兼容：保留原有文件存储功能
'''
    doc.add_paragraph(code_changes)
    
    # 2.4 性能测试结果
    doc.add_heading('2.4 性能测试结果', level=2)
    
    perf_table = doc.add_table(rows=1, cols=3)
    perf_table.style = 'Medium Grid 3 Accent 1'
    hdr_cells = perf_table.rows[0].cells
    headers = ['指标', 'MongoDB', '文件存储']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    perf_data = [
        ['写入QPS', '16,128', '11,953'],
        ['写入延迟', '0.06ms', '0.08ms'],
        ['读取QPS', '761', '30,271'],
        ['读取延迟', '1.31ms', '0.03ms'],
        ['P95延迟', '1.75ms', '0.09ms']
    ]
    
    for item in perf_data:
        row_cells = perf_table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]
    
    doc.add_paragraph('\n📈 关键结论：MongoDB写入性能比文件存储快34.9%，适合高频传感器数据上报场景')
    
    # 2.5 自动化运维
    doc.add_heading('2.5 自动化运维体系', level=2)
    
    ops_info = '''
✅ 自动备份策略
• 定时任务：每日凌晨2点执行
• 备份路径：/backup/mongodb/
• 保留策略：滚动保留最近7天
• 压缩格式：gzip压缩节省空间

✅ 实时监控告警
• 监控频率：每5分钟执行一次
• 监控指标：连接数、内存、磁盘、慢查询
• 告警阈值：
  - 连接数使用率 > 80%
  - 常驻内存 > 2GB
  - 数据量 > 10GB
  - 存在慢查询（>100ms）

✅ 当前监控状态（全部正常）
• 🔗 连接数使用率：0.01% (3/51200)
• 💾 内存使用：125MB
• 💿 磁盘使用：0.45MB
• 🐢 慢查询：0个
'''
    doc.add_paragraph(ops_info)
    
    doc.add_page_break()
    
    # ==================== 三、演示说明 ====================
    doc.add_heading('三、演示说明', level=1)
    
    # 3.1 本地模拟器使用
    doc.add_heading('3.1 本地模拟器使用', level=2)
    
    simulator_guide = '''
📍 模拟器位置：
d:\\学习\\软件设计\\data-server\\examples\\simulator_mq.py

🚀 使用方法：
1. 打开命令行（CMD或PowerShell）
2. 进入项目目录：
   cd d:\\学习\\软件设计\\data-server
3. 运行模拟器：
   python examples/simulator_mq.py

💡 功能说明：
• 模拟生成皮肤传感器数据（moisture, oiliness, temperature）
• 通过Redis消息队列发送数据到服务器
• 支持失败重传机制
• 可配置发送频率和数据量

📊 查看效果：
• 服务器端：python3 scripts/mongodb_monitor.py 查看实时数据增长
• MongoDB：mongo sensor_data -> db.skin_sensor.countDocuments()
'''
    doc.add_paragraph(simulator_guide)
    
    # 3.2 服务器端查看方法
    doc.add_heading('3.2 服务器端查看方法', level=2)
    
    server_guide = '''
🔑 SSH登录服务器：
ssh root@47.103.108.47
密码：@Dierzu999

📊 查看MongoDB数据：
# 方式1：使用mongo shell
mongo sensor_data
db.skin_sensor.find().limit(5).pretty()
db.skin_sensor.countDocuments()

# 方式2：使用Python脚本
cd /root/course-project/week5/data-server/data-server
python3 scripts/mongodb_monitor.py

📁 查看备份文件：
ls -lh /backup/mongodb/

📝 查看监控报告：
cat /var/log/mongodb_monitor_report.json | python3 -m json.tool

🔄 查看定时任务：
crontab -l

📈 手动执行性能测试：
python3 scripts/performance_comparison_test.py 500 50
'''
    doc.add_paragraph(server_guide)
    
    # 3.3 在线演示流程
    doc.add_heading('3.3 在线演示流程建议', level=2)
    
    demo_flow = '''
推荐演示步骤（5-8分钟）：

1️⃣ 展示Git仓库（1分钟）
   • 打开GitHub：https://github.com/MOONFISH2233/software-design-project
   • 切换到week6分支
   • 展示提交记录："第六周任务：MongoDB数据库集成"
   • 展示新增的4个脚本文件

2️⃣ SSH登录服务器（1分钟）
   • 现场SSH连接：ssh root@47.103.108.47
   • 检查MongoDB服务：systemctl status mongod
   • 查看数据库：mongo sensor_data -> show collections

3️⃣ 运行本地模拟器（2分钟）
   • 在本地运行：python examples/simulator_mq.py
   • 同时在服务器端观察：db.skin_sensor.countDocuments()
   • 展示数据实时增长

4️⃣ 展示监控报告（1分钟）
   • 运行：python3 scripts/mongodb_monitor.py
   • 展示各项指标正常
   • 解释监控的意义

5️⃣ 展示备份机制（1分钟）
   • 查看备份目录：ls -lh /backup/mongodb/
   • 手动触发备份：./scripts/mongodb_backup.sh
   • 展示cron配置：crontab -l

6️⃣ 性能对比展示（1分钟）
   • 运行性能测试：python3 scripts/performance_comparison_test.py 500 50
   • 展示结果：MongoDB写入QPS 16,128 vs 文件存储 11,953
   • 强调提升34.9%

💡 演示技巧：
• 提前准备好截图备用（防止网络问题）
• 熟悉每个命令的输出含义
• 准备1-2个常见问题回答
• 控制时间在8分钟以内
'''
    doc.add_paragraph(demo_flow)
    
    doc.add_page_break()
    
    # ==================== 四、飞书上传清单 ====================
    doc.add_heading('四、飞书上传清单', level=1)
    
    upload_list = '''
建议在飞书文档中上传以下内容：

📄 文档类（必传）：
1. ✅ 第六周工作总结与MongoDB使用指南.docx
   - 完整的使用手册
   - 包含所有命令和配置说明
   
2. ✅ 数据库设计说明书.docx
   - ER图、表结构、索引设计
   - 专业的数据库设计文档
   
3. ✅ API接口数据传输参数表.xlsx
   - 12个API接口的详细参数
   - Excel格式便于查阅

📊 报告类（必传）：
4. ✅ 本汇报文档（当前文档）
   - 工作完成情况总结
   - 演示说明和使用指南
   
5. ✅ 性能对比测试报告
   - MongoDB vs 文件存储详细数据
   - 图表展示性能差异

💻 代码类（可选）：
6. Git仓库链接
   - https://github.com/MOONFISH2233/software-design-project
   - 分支：week6
   - 最新提交：3f3f515

🎬 演示类（强烈建议）：
7. 录屏视频（2-3分钟）
   - 录制本地模拟器运行过程
   - 展示服务器端数据实时变化
   - 展示监控报告输出
   
8. 关键截图
   - MongoDB服务运行状态
   - 数据库集合列表
   - 性能测试结果
   - 备份文件列表
   - 监控报告输出

📋 验收清单（必传）：
9. 任务完成情况对照表
   - 列出本周所有要求
   - 标注完成状态（✅/❌）
   - 提供验证方法
'''
    doc.add_paragraph(upload_list)
    
    doc.add_page_break()
    
    # ==================== 五、给老师的展示要点 ====================
    doc.add_heading('五、给老师的展示要点', level=1)
    
    teacher_tips = '''
🎯 核心展示策略：

1️⃣ 结论先行（开场30秒）
"老师好，本周我们完成了MongoDB数据库集成和运维体系建设，包括：
• 安装了MongoDB 6.0.27并设计了3个集合
• 改造了代码支持双写模式和故障转移
• 建立了自动化备份和监控告警系统
• 完成了性能测试，写入QPS提升34.9%"

2️⃣ 证据支撑（2-3分钟）
• 展示Git提交记录证明工作量
• SSH登录服务器现场演示
• 运行监控脚本展示实时状态
• 展示备份文件和cron配置

3️⃣ 技术亮点（1-2分钟）
• 强调高可用设计（故障转移）
• 强调自动化运维（备份+监控）
• 强调性能优化（34.9%提升）
• 强调专业性（完整文档）

4️⃣ 互动演示（2-3分钟）
• 邀请老师提问
• 现场运行模拟器
• 展示数据实时写入
• 解答技术问题

💡 注意事项：
✅ 提前测试SSH连接确保畅通
✅ 准备好备用截图防止网络问题
✅ 熟悉每个命令的作用和输出
✅ 控制总时间在8分钟以内
✅ 重点突出"做了什么"和"怎么验证"

❌ 避免：
❌ 不要照念PPT或文档
❌ 不要陷入技术细节无法自拔
❌ 不要忽略老师的提问
❌ 不要超时影响后续安排

📞 可能的问题及回答：

Q: 为什么选择MongoDB而不是MySQL？
A: 传感器数据是半结构化JSON，MongoDB的文档模型更灵活；写入性能更好（实测快34.9%）；天然支持水平扩展。

Q: 如果MongoDB挂了怎么办？
A: 我们有故障转移机制，会自动降级到文件存储，保证数据不丢失。同时有监控告警，能及时发现并修复。

Q: 备份策略是什么？
A: 每日凌晨2点自动备份，保留7天，gzip压缩。可以手动触发备份，也可以恢复到任意时间点。

Q: 性能瓶颈在哪里？
A: 目前MongoDB写入QPS 16,128，瓶颈在网络IO和磁盘IO。可以通过增加副本集、分片集群进一步提升。

Q: 如何保证数据安全？
A: 三层保障：1) 双写模式冗余存储 2) 每日自动备份 3) 实时监控告警异常。
'''
    doc.add_paragraph(teacher_tips)
    
    doc.add_page_break()
    
    # ==================== 六、快速参考命令 ====================
    doc.add_heading('六、快速参考命令', level=1)
    
    quick_ref = '''
🔧 常用命令速查：

# 1. 本地模拟器
cd d:\\学习\\软件设计\\data-server
python examples/simulator_mq.py

# 2. SSH登录服务器
ssh root@47.103.108.47
# 密码：@Dierzu999

# 3. 查看MongoDB状态
systemctl status mongod

# 4. 连接MongoDB
mongo sensor_data
show collections
db.skin_sensor.countDocuments()
db.skin_sensor.find().limit(5).pretty()

# 5. 运行监控脚本
cd /root/course-project/week5/data-server/data-server
python3 scripts/mongodb_monitor.py

# 6. 手动触发备份
./scripts/mongodb_backup.sh

# 7. 查看备份文件
ls -lh /backup/mongodb/

# 8. 查看定时任务
crontab -l

# 9. 性能测试
python3 scripts/performance_comparison_test.py 500 50

# 10. 查看监控报告
cat /var/log/mongodb_monitor_report.json | python3 -m json.tool

# 11. 查看备份日志
tail -f /var/log/mongodb_backup.log

# 12. 查看监控日志
tail -f /var/log/mongodb_monitor.log

# 13. Git操作
git branch          # 查看分支
git log --oneline   # 查看提交
git pull origin week6  # 拉取最新代码
'''
    doc.add_paragraph(quick_ref)
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(__file__), 'docs', '第六周飞书汇报文档.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ 飞书汇报文档已生成：{output_path}")
    return output_path

if __name__ == '__main__':
    create_feishu_report()
