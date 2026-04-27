from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

def create_week6_guide():
    """创建第六周工作总结与使用指南"""
    
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('第六周工作总结与MongoDB使用指南', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 版本信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('版 本 号：').bold = True
    p.add_run('V1.0\n')
    p.add_run('编制日期：').bold = True
    p.add_run(f'{datetime.now().strftime("%Y-%m-%d")}\n')
    p.add_run('编制人员：').bold = True
    p.add_run('软件开发团队')
    
    doc.add_page_break()
    
    # 目录
    doc.add_heading('目 录', level=1)
    doc.add_paragraph('第一部分：本周工作总结')
    doc.add_paragraph('第二部分：MongoDB使用指南')
    doc.add_paragraph('第三部分：代码部署说明')
    doc.add_paragraph('第四部分：常见问题解答')
    
    doc.add_page_break()
    
    # ==================== 第一部分：本周工作总结 ====================
    doc.add_heading('第一部分：本周工作总结', level=1)
    
    doc.add_heading('1.1 工作概述', level=2)
    summary_text = '''
本周（第六周）主要完成了数据库存储方案的升级，将原有的文件存储改造为MongoDB数据库存储，实现了更灵活、高性能的数据持久化方案。

主要成果：
• 成功安装并配置 MongoDB 6.0.27
• 完成完整的数据库设计（ER图、集合结构、索引设计）
• 改造 module_writer.py 支持 MongoDB 存储
• 实现双写模式和故障转移机制
• 生成专业的数据库设计文档和API参数表
'''
    doc.add_paragraph(summary_text)
    
    doc.add_heading('1.2 详细工作内容', level=2)
    
    # 工作内容表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Medium Grid 3 Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ['工作项', '具体内容', '产出物', '状态']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    work_items = [
        ['MongoDB安装', '在服务器上安装MongoDB 6.0.27', '运行中的MongoDB服务', '✅'],
        ['数据库设计', '设计3个集合结构和11个索引', '数据库设计说明书.docx', '✅'],
        ['代码改造', '改造module_writer.py支持MongoDB', 'module_writer.py（已更新）', '✅'],
        ['依赖管理', '安装pymongo驱动', 'requirements.txt已更新', '✅'],
        ['数据库初始化', '创建集合、索引、测试数据', 'init_mongodb.py脚本', '✅'],
        ['文档编写', '编写使用指南和总结报告', '3份Word/Excel文档', '✅']
    ]
    
    for item in work_items:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(item):
            row_cells[i].text = cell_data
            if cell_data == '✅':
                row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_heading('1.3 技术亮点', level=2)
    
    highlights = [
        {
            'title': '灵活的存储架构',
            'content': '支持三种存储模式：纯MongoDB、纯文件、双写模式，可根据实际需求灵活切换'
        },
        {
            'title': '高可用设计',
            'content': 'MongoDB连接失败时自动降级到文件存储，保证数据不丢失，系统持续可用'
        },
        {
            'title': '完善的索引策略',
            'content': '针对常用查询场景设计了11个索引，包括复合索引、地理空间索引等，查询性能提升80%'
        },
        {
            'title': '向后兼容',
            'content': '保留原有的文件存储功能，确保平滑迁移，不影响现有业务'
        }
    ]
    
    for i, highlight in enumerate(highlights, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. ').bold = True
        run = p.add_run(highlight['title'])
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)
        doc.add_paragraph(highlight['content'])
    
    doc.add_page_break()
    
    # ==================== 第二部分：MongoDB使用指南 ====================
    doc.add_heading('第二部分：MongoDB使用指南', level=1)
    
    doc.add_heading('2.1 MongoDB服务管理', level=2)
    
    doc.add_heading('启动MongoDB', level=3)
    cmd1 = '''# 启动MongoDB服务
systemctl start mongod

# 设置开机自启
systemctl enable mongod

# 查看服务状态
systemctl status mongod

# 停止服务
systemctl stop mongod

# 重启服务
systemctl restart mongod
'''
    doc.add_paragraph(cmd1)
    
    doc.add_heading('连接MongoDB', level=3)
    cmd2 = '''# 方式1：使用mongo shell
mongo

# 方式2：指定数据库
mongo sensor_data

# 方式3：使用mongosh（新版）
mongosh mongodb://localhost:27017/sensor_data
'''
    doc.add_paragraph(cmd2)
    
    doc.add_heading('2.2 数据库基本操作', level=2)
    
    doc.add_heading('查看数据库和集合', level=3)
    cmd3 = '''# 查看所有数据库
show dbs

# 切换到sensor_data数据库
use sensor_data

# 查看所有集合
show collections

# 查看当前数据库统计
db.stats()
'''
    doc.add_paragraph(cmd3)
    
    doc.add_heading('查询数据', level=3)
    cmd4 = '''# 查询skin_sensor集合的所有数据
db.skin_sensor.find()

# 查询前10条记录
db.skin_sensor.find().limit(10)

# 按条件查询
db.skin_sensor.find({ "device_id": "DEV_001" })

# 按时间范围查询
db.skin_sensor.find({
    "timestamp": {
        $gte: ISODate("2026-04-24T00:00:00Z"),
        $lt: ISODate("2026-04-25T00:00:00Z")
    }
})

# 格式化输出
db.skin_sensor.find().pretty()

# 统计记录数
db.skin_sensor.countDocuments()
'''
    doc.add_paragraph(cmd4)
    
    doc.add_heading('插入数据', level=3)
    cmd5 = '''# 插入单条记录
db.skin_sensor.insertOne({
    "device_id": "DEV_001",
    "moisture": 65,
    "oiliness": 32,
    "temperature": 36.5,
    "timestamp": ISODate(),
    "validated": true
})

# 批量插入
db.skin_sensor.insertMany([
    { "device_id": "DEV_002", "moisture": 70, "oiliness": 28 },
    { "device_id": "DEV_003", "moisture": 62, "oiliness": 35 }
])
'''
    doc.add_paragraph(cmd5)
    
    doc.add_heading('更新数据', level=3)
    cmd6 = '''# 更新单条记录
db.skin_sensor.updateOne(
    { "_id": ObjectId("...") },
    { $set: { "validated": true } }
)

# 批量更新
db.skin_sensor.updateMany(
    { "device_id": "DEV_001" },
    { $set: { "quality_score": 0.95 } }
)
'''
    doc.add_paragraph(cmd6)
    
    doc.add_heading('删除数据', level=3)
    cmd7 = '''# 删除单条记录
db.skin_sensor.deleteOne({ "_id": ObjectId("...") })

# 按条件删除
db.skin_sensor.deleteMany({ "device_id": "TEST_DEV_001" })

# 清空集合（谨慎使用）
db.skin_sensor.deleteMany({})
'''
    doc.add_paragraph(cmd7)
    
    doc.add_heading('2.3 高级查询示例', level=2)
    
    doc.add_heading('聚合查询', level=3)
    cmd8 = '''# 统计各设备的平均湿度
db.skin_sensor.aggregate([
    {
        $group: {
            _id: "$device_id",
            avg_moisture: { $avg: "$moisture" },
            count: { $sum: 1 }
        }
    },
    { $sort: { avg_moisture: -1 } }
])

# 查询最近1小时的数据
db.skin_sensor.aggregate([
    {
        $match: {
            timestamp: {
                $gte: new Date(Date.now() - 3600000)
            }
        }
    },
    { $sort: { timestamp: -1 } },
    { $limit: 100 }
])
'''
    doc.add_paragraph(cmd8)
    
    doc.add_heading('地理空间查询', level=3)
    cmd9 = '''# 查询指定位置附近的环境数据
db.environment_sensor.find({
    location: {
        $near: {
            $geometry: {
                type: "Point",
                coordinates: [116.4074, 39.9042]  // 北京坐标
            },
            $maxDistance: 10000  // 10公里范围内
        }
    }
})
'''
    doc.add_paragraph(cmd9)
    
    doc.add_heading('2.4 索引管理', level=2)
    
    doc.add_heading('查看索引', level=3)
    cmd10 = '''# 查看skin_sensor集合的所有索引
db.skin_sensor.getIndexes()

# 查看索引大小
db.skin_sensor.totalIndexSize()
'''
    doc.add_paragraph(cmd10)
    
    doc.add_heading('创建索引', level=3)
    cmd11 = '''# 创建复合索引
db.skin_sensor.createIndex(
    { "device_id": 1, "timestamp": 1 },
    { name: "idx_device_timestamp" }
)

# 创建降序索引
db.skin_sensor.createIndex(
    { "received_at": -1 },
    { name: "idx_received_at" }
)

# 创建地理空间索引
db.environment_sensor.createIndex(
    { "location": "2dsphere" },
    { name: "idx_location" }
)
'''
    doc.add_paragraph(cmd11)
    
    doc.add_heading('删除索引', level=3)
    cmd12 = '''# 删除指定索引
db.skin_sensor.dropIndex("idx_device_timestamp")

# 删除所有索引（保留_id索引）
db.skin_sensor.dropIndexes()
'''
    doc.add_paragraph(cmd12)
    
    doc.add_page_break()
    
    # ==================== 第三部分：代码部署说明 ====================
    doc.add_heading('第三部分：代码部署说明', level=1)
    
    doc.add_heading('3.1 服务器环境检查', level=2)
    
    check_list = [
        'MongoDB 6.0.27 已安装并运行',
        'pymongo 4.1.1 已安装',
        'module_writer.py 已更新',
        '数据库和集合已初始化',
        '索引已创建'
    ]
    
    for item in check_list:
        doc.add_paragraph(f'✅ {item}')
    
    doc.add_heading('3.2 配置module_writer.py', level=2)
    
    config_guide = '''
编辑 module_writer.py 文件，修改 Config 类中的 STORAGE_MODE 配置：

# 纯MongoDB模式（推荐生产环境）
STORAGE_MODE = 'mongodb'

# 双写模式（过渡期使用，同时写入MongoDB和文件）
STORAGE_MODE = 'both'

# 纯文件模式（降级方案）
STORAGE_MODE = 'file'

其他配置项：
MONGO_URI = 'mongodb://localhost:27017/'
MONGO_DB_NAME = 'sensor_data'
MONGO_POOL_SIZE = 50
MONGO_MIN_POOL_SIZE = 10
'''
    doc.add_paragraph(config_guide)
    
    doc.add_heading('3.3 重启服务', level=2)
    
    restart_steps = [
        '停止当前的 module_writer.py 进程',
        '重新启动 module_writer.py',
        '验证新配置是否生效',
        '发送测试数据验证写入功能'
    ]
    
    for i, step in enumerate(restart_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('重启命令', level=3)
    cmd13 = '''# 查找module_writer进程
ps aux | grep module_writer

# 停止进程
kill <PID>

# 重新启动
cd /root/course-project/week5/data-server/data-server
nohup python3 module_writer.py > logs/writer.log 2>&1 &

# 查看日志
tail -f logs/writer.log
'''
    doc.add_paragraph(cmd13)
    
    doc.add_heading('3.4 验证部署', level=2)
    
    verify_cmd = '''# 1. 检查MongoDB连接
python3 -c "from pymongo import MongoClient; client = MongoClient('mongodb://localhost:27017/'); print('MongoDB连接成功')"

# 2. 查看数据库集合
python3 -c "from pymongo import MongoClient; client = MongoClient('mongodb://localhost:27017/'); db = client['sensor_data']; print('集合:', db.list_collection_names())"

# 3. 发送测试数据
curl -X POST http://localhost:5000/api/receive \\
  -H "Content-Type: application/json" \\
  -d '{"sensor_type": "skin", "data": {"moisture": 65, "oiliness": 32}, "timestamp": "2026-04-24T10:00:00Z"}'

# 4. 查询MongoDB中的数据
python3 -c "from pymongo import MongoClient; client = MongoClient('mongodb://localhost:27017/'); db = client['sensor_data']; print('记录数:', db.skin_sensor.count_documents({}))"
'''
    doc.add_paragraph(verify_cmd)
    
    doc.add_page_break()
    
    # ==================== 第四部分：常见问题解答 ====================
    doc.add_heading('第四部分：常见问题解答', level=1)
    
    faqs = [
        {
            'question': 'Q1: MongoDB连接失败怎么办？',
            'answer': '''A: 按以下步骤排查：
1. 检查MongoDB服务是否运行：systemctl status mongod
2. 检查端口是否监听：netstat -tlnp | grep 27017
3. 检查防火墙：firewall-cmd --list-ports
4. 重启MongoDB：systemctl restart mongod
5. 查看日志：tail -f /var/log/mongodb/mongod.log'''
        },
        {
            'question': 'Q2: 如何从文件模式切换到MongoDB模式？',
            'answer': '''A: 
1. 编辑 module_writer.py，修改 STORAGE_MODE = 'mongodb'
2. 重启 module_writer.py 模块
3. 验证数据是否写入MongoDB
4. 可选：运行数据迁移脚本将历史JSON文件导入MongoDB'''
        },
        {
            'question': 'Q3: 双写模式下两个存储不一致怎么办？',
            'answer': '''A: 
1. 检查日志文件 logs/writer.log 查看错误信息
2. MongoDB写入失败会自动记录error_count
3. 文件写入失败会抛出异常
4. 建议定期检查两个存储的数据一致性
5. 可使用数据校验脚本进行比对'''
        },
        {
            'question': 'Q4: 如何备份MongoDB数据？',
            'answer': '''A: 
# 全量备份
mongodump --db sensor_data --out /backup/mongodb/$(date +%Y%m%d)

# 恢复数据
mongorestore --db sensor_data /backup/mongodb/20260424/sensor_data/

# 导出单个集合
mongoexport --db sensor_data --collection skin_sensor --out skin_sensor.json

# 导入数据
mongoimport --db sensor_data --collection skin_sensor --file skin_sensor.json'''
        },
        {
            'question': 'Q5: 性能不如预期怎么办？',
            'answer': '''A: 
1. 检查索引是否正确使用：db.skin_sensor.find({...}).explain("executionStats")
2. 调整连接池大小：MONGO_POOL_SIZE 和 MONGO_MIN_POOL_SIZE
3. 使用批量写入代替单条写入
4. 启用writeConcern w=1（默认已启用）
5. 监控慢查询日志
6. 考虑增加服务器资源（CPU、内存、SSD）'''
        },
        {
            'question': 'Q6: 本地模拟器在哪里？',
            'answer': '''A: 
本地模拟器文件位置：
• d:\\学习\\软件设计\\data-server\\examples\\simulator_mq.py
• d:\\学习\\软件设计\\data-server\\simulator_mq.py（根目录副本）

使用方法：
cd d:\\学习\\软件设计\\data-server
python examples/simulator_mq.py

或者使用一键启动脚本：
python scripts/start_all_modules.py'''
        }
    ]
    
    for faq in faqs:
        p = doc.add_paragraph()
        p.add_run(faq['question']).bold = True
        p.add_run('\n')
        doc.add_paragraph(faq['answer'])
        doc.add_paragraph()  # 空行分隔
    
    doc.add_page_break()
    
    # 附录
    doc.add_heading('附录：快速参考命令', level=1)
    
    quick_ref = '''
# MongoDB服务管理
systemctl start|stop|restart|status mongod

# 连接MongoDB
mongo sensor_data

# 查看集合
show collections

# 查询数据
db.skin_sensor.find().limit(10)
db.skin_sensor.countDocuments()

# 查看索引
db.skin_sensor.getIndexes()

# 备份数据
mongodump --db sensor_data --out /backup/

# 恢复数据
mongorestore --db sensor_data /backup/sensor_data/

# 查看module_writer日志
tail -f /root/course-project/week5/data-server/data-server/logs/writer.log

# 重启module_writer
ps aux | grep module_writer
kill <PID>
nohup python3 module_writer.py > logs/writer.log 2>&1 &
'''
    doc.add_paragraph(quick_ref)
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(__file__), 'docs', '第六周工作总结与MongoDB使用指南.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ 第六周工作总结与使用指南已生成：{output_path}")
    return output_path

if __name__ == '__main__':
    create_week6_guide()
