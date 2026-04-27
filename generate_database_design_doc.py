from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_database_design_doc():
    """创建数据库设计说明书Word文档"""
    
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('传感器数据服务器 - 数据库设计说明书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 版本信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('版 本 号：').bold = True
    p.add_run('V1.0\n')
    p.add_run('编制日期：').bold = True
    p.add_run('2026-04-24\n')
    p.add_run('编制人员：').bold = True
    p.add_run('软件开发团队')
    
    doc.add_page_break()
    
    # 目录
    doc.add_heading('目 录', level=1)
    doc.add_paragraph('1. 项目概述')
    doc.add_paragraph('2. 数据库选型说明')
    doc.add_paragraph('3. 概念模型设计（ER图）')
    doc.add_paragraph('4. 逻辑模型设计（集合结构）')
    doc.add_paragraph('5. 物理模型设计（索引优化）')
    doc.add_paragraph('6. 数据流设计')
    doc.add_paragraph('7. 实施计划')
    
    doc.add_page_break()
    
    # 1. 项目概述
    doc.add_heading('1. 项目概述', level=1)
    
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph(
        '本项目为传感器数据服务器，负责接收来自多个IoT设备上传的传感器数据，'
        '包括皮肤传感器、环境传感器和设备状态数据。系统采用消息队列架构实现高并发处理，'
        '需要设计合理的数据库存储方案以支持数据的持久化存储和高效查询。'
    )
    
    doc.add_heading('1.2 设计目标', level=2)
    p = doc.add_paragraph()
    p.add_run('• ').bold = True
    p.add_run('支持高并发写入：QPS ≥ 500\n')
    p.add_run('• ').bold = True
    p.add_run('数据存储可靠性：99.99%\n')
    p.add_run('• ').bold = True
    p.add_run('查询响应时间：< 100ms\n')
    p.add_run('• ').bold = True
    p.add_run('支持水平扩展\n')
    p.add_run('• ').bold = True
    p.add_run('数据可追溯性：保留至少6个月历史数据')
    
    doc.add_heading('1.3 数据类型', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ['数据类型', '字段数量', '写入频率', '存储周期']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    data_rows = [
        ['皮肤传感器', '8个', '2秒/次', '6个月'],
        ['环境传感器', '10个', '3秒/次', '6个月'],
        ['设备状态', '7个', '30秒/次', '1年']
    ]
    
    for row_data in data_rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    doc.add_page_break()
    
    # 2. 数据库选型说明
    doc.add_heading('2. 数据库选型说明', level=1)
    
    doc.add_heading('2.1 选型对比', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Medium Grid 3 Accent 1'
    
    hdr_cells = table.rows[0].cells
    headers = ['对比项', 'MongoDB', 'MySQL', 'PostgreSQL', 'Redis']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    comparison_data = [
        ['数据结构', '文档型（灵活）', '关系型（固定）', '关系型（固定）', '键值对'],
        ['写入性能', '★★★★★', '★★★☆☆', '★★★☆☆', '★★★★★'],
        ['查询灵活性', '★★★★☆', '★★★★★', '★★★★★', '★★☆☆☆'],
        ['水平扩展', '原生支持', '需分库分表', '需分库分表', '集群模式'],
        ['Schema变更', '无需迁移', '需ALTER TABLE', '需ALTER TABLE', '不适用'],
        ['最终选择', '✅ 推荐', '备选', '备选', '缓存层']
    ]
    
    for row_data in comparison_data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    doc.add_heading('2.2 选择MongoDB的理由', level=2)
    p = doc.add_paragraph()
    p.add_run('1. 灵活的文档模型：').bold = True
    p.add_run('传感器数据字段可能变化，MongoDB的无Schema特性便于扩展\n')
    
    p = doc.add_paragraph()
    p.add_run('2. 高性能写入：').bold = True
    p.add_run('基于BSON的二进制格式，写入性能优于传统关系型数据库\n')
    
    p = doc.add_paragraph()
    p.add_run('3. 水平扩展能力：').bold = True
    p.add_run('原生支持分片（Sharding），可轻松应对数据量增长\n')
    
    p = doc.add_paragraph()
    p.add_run('4. 丰富的查询功能：').bold = True
    p.add_run('支持聚合管道、地理空间查询等高级功能\n')
    
    p = doc.add_paragraph()
    p.add_run('5. 生态系统完善：').bold = True
    p.add_run('Python驱动pymongo成熟稳定，社区活跃')
    
    doc.add_page_break()
    
    # 3. 概念模型设计
    doc.add_heading('3. 概念模型设计（ER图）', level=1)
    
    doc.add_heading('3.1 实体定义', level=2)
    
    doc.add_heading('实体1：Device（设备）', level=3)
    p = doc.add_paragraph()
    p.add_run('属性：').bold = True
    p.add_run('device_id（主键）、device_type、firmware_version、install_date、location\n')
    p.add_run('说明：').bold = True
    p.add_run('表示一个物理传感器设备')
    
    doc.add_heading('实体2：SkinSensorData（皮肤传感器数据）', level=3)
    p = doc.add_paragraph()
    p.add_run('属性：').bold = True
    p.add_run('_id（主键）、device_id（外键）、moisture、oiliness、temperature、timestamp\n')
    p.add_run('说明：').bold = True
    p.add_run('记录皮肤传感器的测量数据')
    
    doc.add_heading('实体3：EnvironmentSensorData（环境传感器数据）', level=3)
    p = doc.add_paragraph()
    p.add_run('属性：').bold = True
    p.add_run('_id（主键）、device_id（外键）、temperature、humidity、pm25、co2、timestamp\n')
    p.add_run('说明：').bold = True
    p.add_run('记录环境传感器的测量数据')
    
    doc.add_heading('实体4：DeviceStatus（设备状态）', level=3)
    p = doc.add_paragraph()
    p.add_run('属性：').bold = True
    p.add_run('_id（主键）、device_id（外键）、status、battery_level、signal_strength、last_heartbeat\n')
    p.add_run('说明：').bold = True
    p.add_run('记录设备的运行状态')
    
    doc.add_heading('3.2 实体关系', level=2)
    p = doc.add_paragraph()
    p.add_run('• Device 1 : N SkinSensorData（一对多）\n')
    p.add_run('• Device 1 : N EnvironmentSensorData（一对多）\n')
    p.add_run('• Device 1 : N DeviceStatus（一对多）')
    
    doc.add_heading('3.3 ER图示意', level=2)
    er_diagram = '''
┌─────────────┐       ┌──────────────────┐
│   Device    │1     N│ SkinSensorData   │
├─────────────┤───────├──────────────────┤
│ device_id   │       │ _id              │
│ device_type │       │ device_id (FK)   │
│ firmware    │       │ moisture         │
│ location    │       │ oiliness         │
└─────────────┘       │ temperature      │
                      │ timestamp        │
                      └──────────────────┘
           1│                    ▲
            │                    │
            │                    │N
            │          ┌──────────────────┐
            └──────────│EnvironmentSensor │
                       ├──────────────────┤
                       │ _id              │
                       │ device_id (FK)   │
                       │ temperature      │
                       │ humidity         │
                       │ pm25             │
                       │ co2              │
                       │ timestamp        │
                       └──────────────────┘
            
            1│                    ▲
             │                    │
             │                    │N
             │          ┌──────────────────┐
             └──────────│  DeviceStatus    │
                        ├──────────────────┤
                        │ _id              │
                        │ device_id (FK)   │
                        │ status           │
                        │ battery_level    │
                        │ signal_strength  │
                        │ last_heartbeat   │
                        └──────────────────┘
    '''
    doc.add_paragraph(er_diagram)
    
    doc.add_page_break()
    
    # 4. 逻辑模型设计
    doc.add_heading('4. 逻辑模型设计（集合结构）', level=1)
    
    doc.add_heading('4.1 skin_sensor 集合', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['字段名', '类型', '必填', '说明', '示例']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    skin_fields = [
        ['_id', 'ObjectId', '是', '主键，自动生成', '507f1f77bcf86cd799439011'],
        ['device_id', 'String', '是', '设备唯一标识', 'DEV_001'],
        ['moisture', 'Int32', '是', '皮肤湿度（0-100）', '65'],
        ['oiliness', 'Int32', '是', '皮肤油脂度（0-100）', '32'],
        ['temperature', 'Double', '否', '皮肤温度（℃）', '36.5'],
        ['timestamp', 'ISODate', '是', '数据采集时间', '2026-04-24T08:30:00Z'],
        ['received_at', 'ISODate', '是', '服务器接收时间', '2026-04-24T08:30:01Z'],
        ['client_ip', 'String', '否', '客户端IP地址', '192.168.1.100'],
        ['request_id', 'String', '否', '请求追踪ID', 'req_abc123'],
        ['validated', 'Boolean', '是', '是否通过验证', 'true'],
        ['quality_score', 'Double', '否', '数据质量评分（0-1）', '0.95']
    ]
    
    for field in skin_fields:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(field):
            row_cells[i].text = str(cell_data)
    
    doc.add_heading('4.2 environment_sensor 集合', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    env_fields = [
        ['_id', 'ObjectId', '是', '主键', '507f1f77bcf86cd799439012'],
        ['device_id', 'String', '是', '设备唯一标识', 'DEV_002'],
        ['temperature', 'Double', '是', '环境温度（℃）', '25.3'],
        ['humidity', 'Double', '是', '环境湿度（%）', '60.5'],
        ['pm25', 'Int32', '是', 'PM2.5浓度（μg/m³）', '35'],
        ['co2', 'Int32', '是', 'CO₂浓度（ppm）', '450'],
        ['light_intensity', 'Double', '否', '光照强度（lux）', '500.0'],
        ['noise_level', 'Double', '否', '噪音等级（dB）', '45.2'],
        ['timestamp', 'ISODate', '是', '数据采集时间', '2026-04-24T08:30:00Z'],
        ['received_at', 'ISODate', '是', '服务器接收时间', '2026-04-24T08:30:01Z'],
        ['location', 'Object', '否', '地理位置', '{lat: 39.9, lng: 116.4}'],
        ['location.latitude', 'Double', '否', '纬度', '39.9042'],
        ['location.longitude', 'Double', '否', '经度', '116.4074']
    ]
    
    for field in env_fields:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(field):
            row_cells[i].text = str(cell_data)
    
    doc.add_heading('4.3 device_status 集合', level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    status_fields = [
        ['_id', 'ObjectId', '是', '主键', '507f1f77bcf86cd799439013'],
        ['device_id', 'String', '是', '设备唯一标识', 'DEV_001'],
        ['status', 'String', '是', '设备状态', 'online'],
        ['battery_level', 'Int32', '是', '电池电量（0-100）', '85'],
        ['signal_strength', 'Int32', '是', '信号强度（-100~0 dBm）', '-65'],
        ['firmware_version', 'String', '是', '固件版本', 'v2.1.0'],
        ['last_heartbeat', 'ISODate', '是', '最后心跳时间', '2026-04-24T08:35:00Z'],
        ['error_code', 'String', '否', '错误代码', 'ERR_TIMEOUT'],
        ['uptime', 'Int64', '否', '运行时长（秒）', '86400'],
        ['memory_usage', 'Double', '否', '内存使用率（%）', '45.2']
    ]
    
    for field in status_fields:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(field):
            row_cells[i].text = str(cell_data)
    
    doc.add_page_break()
    
    # 5. 物理模型设计
    doc.add_heading('5. 物理模型设计（索引优化）', level=1)
    
    doc.add_heading('5.1 skin_sensor 索引设计', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    idx_headers = ['索引名称', '索引字段', '索引类型', '用途说明']
    for i, header in enumerate(idx_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    skin_indexes = [
        ['idx_device_timestamp', 'device_id + timestamp', '复合升序', '按设备查询某时间段数据'],
        ['idx_received_at', 'received_at', '单字段降序', '按接收时间倒序查询最新数据'],
        ['idx_quality_score', 'quality_score', '单字段降序', '筛选高质量数据'],
        ['idx_validated', 'validated', '单字段', '快速筛选已验证数据']
    ]
    
    for idx in skin_indexes:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(idx):
            row_cells[i].text = cell_data
    
    doc.add_heading('5.2 environment_sensor 索引设计', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(idx_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    env_indexes = [
        ['idx_device_timestamp', 'device_id + timestamp', '复合升序', '按设备查询时间段数据'],
        ['idx_location', 'location', '2dsphere', '地理空间查询（附近设备）'],
        ['idx_pm25', 'pm25', '单字段升序', '空气质量筛选'],
        ['idx_temperature_range', 'temperature', '单字段', '温度范围查询']
    ]
    
    for idx in env_indexes:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(idx):
            row_cells[i].text = cell_data
    
    doc.add_heading('5.3 device_status 索引设计', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(idx_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    status_indexes = [
        ['idx_device_last_heartbeat', 'device_id + last_heartbeat', '复合降序', '查询设备最新状态'],
        ['idx_status', 'status', '单字段', '按状态筛选（在线/离线）'],
        ['idx_battery_low', 'battery_level', '单字段升序', '查找低电量设备']
    ]
    
    for idx in status_indexes:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(idx):
            row_cells[i].text = cell_data
    
    doc.add_heading('5.4 索引创建脚本', level=2)
    script_text = '''// MongoDB Shell 索引创建命令

// 1. skin_sensor 集合索引
db.skin_sensor.createIndex({ "device_id": 1, "timestamp": 1 })
db.skin_sensor.createIndex({ "received_at": -1 })
db.skin_sensor.createIndex({ "quality_score": -1 })
db.skin_sensor.createIndex({ "validated": 1 })

// 2. environment_sensor 集合索引
db.environment_sensor.createIndex({ "device_id": 1, "timestamp": 1 })
db.environment_sensor.createIndex({ "location": "2dsphere" })
db.environment_sensor.createIndex({ "pm25": 1 })
db.environment_sensor.createIndex({ "temperature": 1 })

// 3. device_status 集合索引
db.device_status.createIndex({ "device_id": 1, "last_heartbeat": -1 })
db.device_status.createIndex({ "status": 1 })
db.device_status.createIndex({ "battery_level": 1 })
'''
    doc.add_paragraph(script_text)
    
    doc.add_page_break()
    
    # 6. 数据流设计
    doc.add_heading('6. 数据流设计', level=1)
    
    doc.add_heading('6.1 数据写入流程', level=2)
    flow_text = '''
┌──────────────┐     ┌──────────────┐     ┌──────────────┐
│ IoT 设备     │────▶│ Redis Stream │────▶│ Flask API    │
│ (传感器)     │ HTTP │ sensor:raw   │ MQ   │ /api/receive │
└──────────────┘     └──────────────┘     └──────┬───────┘
                                                  │
                                                  ▼
                                         ┌──────────────┐
                                         │module_writer │
                                         │ (消费MQ)     │
                                         └──────┬───────┘
                                                │
                                    ┌───────────┴───────────┐
                                    ▼                       ▼
                           ┌──────────────┐        ┌──────────────┐
                           │  MongoDB     │        │ JSON 文件    │
                           │ (主存储)     │        │ (备份)       │
                           └──────────────┘        └──────────────┘
'''
    doc.add_paragraph(flow_text)
    
    doc.add_heading('6.2 数据读取流程', level=2)
    read_flow = '''
┌──────────────┐     ┌──────────────┐     ┌──────────────┐
│ 前端应用     │────▶│ Flask API    │────▶│ MongoDB      │
│ / Web界面    │ HTTP │ /api/query   │ Query│ Collection   │
└──────────────┘     └──────────────┘     └──────────────┘
                            │
                            ▼
                   ┌─────────────────┐
                   │ 聚合管道处理     │
                   │ - 过滤           │
                   │ - 分组           │
                   │ - 排序           │
                   └─────────────────┘
'''
    doc.add_paragraph(read_flow)
    
    doc.add_heading('6.3 典型查询场景', level=2)
    
    doc.add_heading('场景1：查询某设备最近1小时的皮肤数据', level=3)
    query1 = '''
db.skin_sensor.find({
    device_id: "DEV_001",
    timestamp: {
        $gte: new Date(Date.now() - 3600000)
    }
}).sort({ timestamp: -1 })
'''
    doc.add_paragraph(query1)
    
    doc.add_heading('场景2：统计各设备的平均PM2.5浓度', level=3)
    query2 = '''
db.environment_sensor.aggregate([
    {
        $group: {
            _id: "$device_id",
            avg_pm25: { $avg: "$pm25" },
            count: { $sum: 1 }
        }
    },
    { $sort: { avg_pm25: -1 } }
])
'''
    doc.add_paragraph(query2)
    
    doc.add_heading('场景3：查找所有离线设备', level=3)
    query3 = '''
db.device_status.find({
    status: "offline",
    last_heartbeat: {
        $lt: new Date(Date.now() - 300000)  // 5分钟未心跳
    }
})
'''
    doc.add_paragraph(query3)
    
    doc.add_page_break()
    
    # 7. 实施计划
    doc.add_heading('7. 实施计划', level=1)
    
    doc.add_heading('7.1 部署步骤', level=2)
    steps = [
        '1. 安装MongoDB 6.0+',
        '   - yum install -y mongodb-org',
        '   - systemctl start mongod',
        '',
        '2. 安装Python依赖',
        '   - pip install pymongo==4.6.1',
        '',
        '3. 初始化数据库',
        '   - 执行 scripts/init_mongodb.py',
        '   - 创建集合和索引',
        '',
        '4. 配置module_writer.py',
        '   - 设置 storage_mode = "mongodb"',
        '   - 配置连接字符串',
        '',
        '5. 重启服务',
        '   - systemctl restart gunicorn-flask-data-server',
        '   - 重启 module_writer.py 模块',
        '',
        '6. 验证测试',
        '   - 发送测试数据',
        '   - 检查MongoDB中的数据',
        '   - 性能压测对比'
    ]
    for step in steps:
        doc.add_paragraph(step)
    
    doc.add_heading('7.2 数据迁移方案', level=2)
    migration = '''
从JSON文件迁移到MongoDB的步骤：

1. 编写迁移脚本 migrate_json_to_mongo.py
2. 遍历 data/ 目录下所有JSON文件
3. 解析JSON数据并批量插入MongoDB
4. 验证数据完整性（记录数对比）
5. 确认无误后删除旧JSON文件（可选）

预计迁移时间：
- 10万条数据：约5-10分钟
- 100万条数据：约1-2小时
'''
    doc.add_paragraph(migration)
    
    doc.add_heading('7.3 备份策略', level=2)
    backup_plan = '''
自动备份方案：

1. 每日凌晨2点执行全量备份
   mongodump --db sensor_data --out /backup/mongodb/$(date +%Y%m%d)

2. 每小时执行增量备份（Oplog）
   
3. 保留策略：
   - 每日备份保留7天
   - 每周备份保留4周
   - 每月备份保留12个月

4. 异地备份：
   - 同步备份到阿里云OSS
   - 成本：约¥50/月（100GB）
'''
    doc.add_paragraph(backup_plan)
    
    doc.add_page_break()
    
    # 附录
    doc.add_heading('附录A：MongoDB常用命令速查', level=1)
    
    commands = '''
# 连接MongoDB
mongo mongodb://localhost:27017

# 查看数据库
show dbs

# 切换数据库
use sensor_data

# 查看集合
show collections

# 查询数据
db.skin_sensor.find().limit(10)
db.skin_sensor.countDocuments()

# 插入数据
db.skin_sensor.insertOne({...})
db.skin_sensor.insertMany([{...}, {...}])

# 更新数据
db.skin_sensor.updateOne(
    { _id: ObjectId("...") },
    { $set: { validated: true } }
)

# 删除数据
db.skin_sensor.deleteOne({ _id: ObjectId("...") })

# 查看索引
db.skin_sensor.getIndexes()

# 性能分析
db.skin_sensor.find({ device_id: "DEV_001" }).explain("executionStats")

# 备份数据库
mongodump --db sensor_data --out /backup/

# 恢复数据库
mongorestore --db sensor_data /backup/sensor_data/
'''
    doc.add_paragraph(commands)
    
    doc.add_heading('附录B：性能优化建议', level=1)
    optimization = '''
1. 连接池优化
   - maxPoolSize: 50-100
   - minPoolSize: 10
   - maxIdleTimeMS: 30000

2. 写入优化
   - 使用批量写入（insertMany）
   - 设置 ordered=false 提高容错
   - 启用writeConcern w=1（默认）

3. 查询优化
   - 避免全表扫描，确保使用索引
   - 限制返回字段（projection）
   - 使用分页（skip+limit 或 cursor）

4. 监控指标
   - 慢查询日志（>100ms）
   - 连接数监控
   - 内存使用情况
   - 磁盘I/O

5. 硬件建议
   - CPU: 4核以上
   - 内存: 8GB以上
   - 磁盘: SSD（IOPS > 3000）
   - 网络: 千兆以太网
'''
    doc.add_paragraph(optimization)
    
    # 保存文档
    output_path = os.path.join(os.path.dirname(__file__), 'docs', '数据库设计说明书.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    
    print(f"✅ 数据库设计说明书已生成：{output_path}")
    return output_path

if __name__ == '__main__':
    create_database_design_doc()
