#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成第二阶段验收快速操作指南 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_quick_guide():
    """创建快速操作指南"""
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('第二阶段验收 - 快速操作指南', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(33, 150, 243)
    
    # 基本信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('服务器: 47.103.108.47\n').bold = True
    info_para.add_run('SSH账号: root / @Dierzu999\n')
    info_para.add_run('服务端口: 5000\n')
    info_para.add_run('验收日期: 2026年4月17日')
    
    doc.add_page_break()
    
    # ========== 任务1 ==========
    h1 = doc.add_heading('任务(1): RESTful API接口演示', level=1)
    h1.runs[0].font.color.rgb = RGBColor(33, 150, 243)
    
    doc.add_paragraph('步骤1: SSH登录服务器', style='Heading 3')
    p1 = doc.add_paragraph('ssh root@47.103.108.47')
    p1.style = 'No Spacing'
    p1.runs[0].font.name = 'Consolas'
    p1.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('步骤2: 查看服务进程', style='Heading 3')
    p2 = doc.add_paragraph('ps aux | grep gunicorn')
    p2.style = 'No Spacing'
    p2.runs[0].font.name = 'Consolas'
    p2.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('预期输出: 1个master + 8个worker进程')
    
    doc.add_paragraph('步骤3: 健康检查', style='Heading 3')
    p3 = doc.add_paragraph('curl http://localhost:5000/api/health | python3 -m json.tool')
    p3.style = 'No Spacing'
    p3.runs[0].font.name = 'Consolas'
    p3.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('汇报要点:', style='Heading 3').runs[0].bold = True
    doc.add_paragraph('• 12个RESTful API接口', style='List Bullet')
    doc.add_paragraph('• Gunicorn 8 workers运行正常', style='List Bullet')
    doc.add_paragraph('• 所有接口遵循RESTful规范', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 任务2 ==========
    h2 = doc.add_heading('任务(2): 数据上传与保存演示', level=1)
    h2.runs[0].font.color.rgb = RGBColor(33, 150, 243)
    
    doc.add_paragraph('步骤1: 发送测试数据', style='Heading 3')
    p4 = doc.add_paragraph('''curl -X POST http://localhost:5000/api/receive \\
  -H "Content-Type: application/json" \\
  -d '{"sensor_type":"skin","moisture":65,"oiliness":32,"timestamp":"2026-04-17 11:00:00"}' ''')
    p4.style = 'No Spacing'
    p4.runs[0].font.name = 'Consolas'
    p4.runs[0].font.size = Pt(9)
    
    doc.add_paragraph('步骤2: 验证文件已创建', style='Heading 3')
    p5 = doc.add_paragraph('ls -lh /root/data-server/data/skin/ | tail -3')
    p5.style = 'No Spacing'
    p5.runs[0].font.name = 'Consolas'
    p5.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('步骤3: 统计数据总量', style='Heading 3')
    p6 = doc.add_paragraph('find /root/data-server/data/ -name "*.json" | wc -l')
    p6.style = 'No Spacing'
    p6.runs[0].font.name = 'Consolas'
    p6.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('汇报要点:', style='Heading 3').runs[0].bold = True
    doc.add_paragraph('• 15000+数据文件已存储', style='List Bullet')
    doc.add_paragraph('• JSON格式，按类型分目录', style='List Bullet')
    doc.add_paragraph('• 处理时间<15ms，性能优秀', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 任务5 ==========
    h5 = doc.add_heading('任务(5): 鉴权功能演示', level=1)
    h5.runs[0].font.color.rgb = RGBColor(33, 150, 243)
    
    doc.add_paragraph('演示1: 无效Token被拒绝', style='Heading 3')
    p7 = doc.add_paragraph('curl -X POST http://localhost:5000/api/receive -H "Authorization: Bearer invalid_token" -d \'{}\'')
    p7.style = 'No Spacing'
    p7.runs[0].font.name = 'Consolas'
    p7.runs[0].font.size = Pt(9)
    doc.add_paragraph('预期输出: 401 Unauthorized')
    
    doc.add_paragraph('演示2: 获取有效Token', style='Heading 3')
    p8 = doc.add_paragraph('curl -X POST http://localhost:5000/api/auth/login -d \'{"username":"admin","password":"admin123"}\'')
    p8.style = 'No Spacing'
    p8.runs[0].font.name = 'Consolas'
    p8.runs[0].font.size = Pt(9)
    
    doc.add_paragraph('演示3: 使用Token成功访问', style='Heading 3')
    p9 = doc.add_paragraph('curl -X POST http://localhost:5000/api/receive -H "Authorization: Bearer <token>" -d \'{}\'')
    p9.style = 'No Spacing'
    p9.runs[0].font.name = 'Consolas'
    p9.runs[0].font.size = Pt(9)
    
    doc.add_paragraph('汇报要点:', style='Heading 3').runs[0].bold = True
    doc.add_paragraph('• JWT Token认证（24小时有效期）', style='List Bullet')
    doc.add_paragraph('• API Key认证（适合IoT设备）', style='List Bullet')
    doc.add_paragraph('• 速率限制：50次/分钟', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 任务7 ==========
    h7 = doc.add_heading('任务(7): 消息队列演示', level=1)
    h7.runs[0].font.color.rgb = RGBColor(33, 150, 243)
    
    doc.add_paragraph('步骤1: 检查Redis服务', style='Heading 3')
    p10 = doc.add_paragraph('redis-cli ping')
    p10.style = 'No Spacing'
    p10.runs[0].font.name = 'Consolas'
    p10.runs[0].font.size = Pt(10)
    doc.add_paragraph('预期输出: PONG')
    
    doc.add_paragraph('步骤2: 查看Streams', style='Heading 3')
    p11 = doc.add_paragraph('redis-cli keys \'sensor:*\'')
    p11.style = 'No Spacing'
    p11.runs[0].font.name = 'Consolas'
    p11.runs[0].font.size = Pt(10)
    doc.add_paragraph('预期输出: sensor:raw, sensor:validated, sensor:write, sensor:logs')
    
    doc.add_paragraph('步骤3: 启动模拟器', style='Heading 3')
    p12 = doc.add_paragraph('cd /root/data-server && python3 simulator_mq.py')
    p12.style = 'No Spacing'
    p12.runs[0].font.name = 'Consolas'
    p12.runs[0].font.size = Pt(10)
    doc.add_paragraph('运行30秒后按Ctrl+C停止')
    
    doc.add_paragraph('汇报要点:', style='Heading 3').runs[0].bold = True
    doc.add_paragraph('• Redis Streams实现消息队列', style='List Bullet')
    doc.add_paragraph('• 4个Stream分类存储', style='List Bullet')
    doc.add_paragraph('• 解耦架构，支持水平扩展', style='List Bullet')
    
    doc.add_page_break()
    
    # ========== 答辩问题 ==========
    h_qa = doc.add_heading('答辩问题速查表', level=1)
    h_qa.runs[0].font.color.rgb = RGBColor(76, 175, 80)
    
    questions = [
        ('Q1: 为什么选择Flask而不是Django？', 
         'A: Flask轻量灵活，适合RESTful API。配合Gunicorn可实现高并发。Django更适合大型Web应用，对于纯API服务过于重量级。'),
        ('Q2: Redis Streams和传统消息队列有什么区别？',
         'A: Redis Streams轻量、无需额外部署、学习成本低。我们已有Redis基础设施，复用降低成本。RabbitMQ/Kafka适合企业级应用，但部署复杂。'),
        ('Q3: 如何保证数据不丢失？',
         'A: 三重保障：①同步写入JSON文件立即落盘；②Redis Stream持久化；③备份机制。即使服务崩溃，数据也能恢复。'),
        ('Q4: 安全性如何评估？',
         'A: 四层防护：认证层（JWT/API Key）、传输层（AES加密）、应用层（速率限制）、审计层（日志记录）。已通过多项安全测试。'),
        ('Q5: 性能瓶颈在哪里？如何优化？',
         'A: 当前瓶颈在磁盘I/O。优化方案：批量写入、SSD硬盘、数据库替代文件存储、异步I/O。理论可支撑10000+ QPS。'),
        ('Q6: 如果要多机部署怎么办？',
         'A: Nginx负载均衡 + Redis集群 + 共享存储/NFS + 数据库读写分离。架构天然支持水平扩展。'),
    ]
    
    for q, a in questions:
        para_q = doc.add_paragraph()
        para_q.add_run(q).bold = True
        para_q.add_run('\n' + a)
        para_q.paragraph_format.space_after = Pt(10)
    
    doc.add_page_break()
    
    # ========== 应急处理 ==========
    h_emergency = doc.add_heading('应急处理方案', level=1)
    h_emergency.runs[0].font.color.rgb = RGBColor(255, 152, 0)
    
    emergencies = [
        ('SSH连接失败', '检查网络：ping 47.103.108.47；检查阿里云安全组；使用备用截图'),
        ('Gunicorn未运行', '重启：cd /root/data-server && nohup gunicorn -c config/gunicorn_config.py app:app &'),
        ('Redis未运行', '启动：systemctl start redis'),
        ('依赖缺失', '安装：pip3 install flask-limiter cryptography redis'),
        ('PPT无法打开', '使用WPS Office；或转换为PDF'),
    ]
    
    for issue, solution in emergencies:
        para = doc.add_paragraph()
        para.add_run(f'问题: {issue}\n').bold = True
        para.add_run(f'解决: {solution}')
        para.paragraph_format.space_after = Pt(8)
    
    # 页脚
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('祝验收顺利！🚀').font.size = Pt(10)
    footer_para.add_run('\n文档版本: v1.0 | 最后更新: 2026年4月17日').font.size = Pt(8)
    
    # 保存
    output_file = 'd:\\学习\\软件设计\\data-server\\docs\\验收快速操作指南.docx'
    doc.save(output_file)
    print(f'✅ Word文档已生成: {output_file}')
    print(f'📄 页数: {len(doc.paragraphs)} 段落')
    
    return output_file

if __name__ == '__main__':
    try:
        create_quick_guide()
    except Exception as e:
        print(f'❌ 生成失败: {e}')
        import traceback
        traceback.print_exc()
