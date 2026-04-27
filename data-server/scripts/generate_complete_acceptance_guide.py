#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成第二阶段完整验收指南 Word 文档
包含：服务启动、鉴权机制、协议文档、本地模拟器演示
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_complete_guide():
    doc = Document()
    
    # 标题
    title = doc.add_heading('Flask 数据服务器 - 第二阶段完整验收指南', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('包含服务启动、鉴权机制、协议文档、本地模拟器演示')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    # 目录
    doc.add_heading('目录', level=1)
    toc_items = [
        '一、服务启动与状态检查',
        '二、鉴权机制详解',
        '三、协议文档说明',
        '四、本地模拟器演示流程',
        '五、完整验收演示脚本',
        '六、常见问题与解决方案'
    ]
    for i, item in enumerate(toc_items, 1):
        p = doc.add_paragraph()
        p.add_run(item).bold = True
    
    doc.add_page_break()
    
    # 第一部分：服务启动
    doc.add_heading('一、服务启动与状态检查', level=1)
    
    doc.add_heading('1.1 问题：Connection refused', level=2)
    p = doc.add_paragraph()
    p.add_run('错误信息：').bold = True
    doc.add_paragraph('curl: (7) Failed to connect to localhost port 5000: Connection refused')
    
    p = doc.add_paragraph()
    p.add_run('原因：').bold = True
    doc.add_paragraph('Gunicorn 服务未启动或已停止')
    
    doc.add_heading('1.2 解决方案：启动服务', level=2)
    
    doc.add_paragraph('方法 1：使用 systemctl 启动（推荐）', style='Heading 3')
    code1 = doc.add_paragraph()
    code1.add_run('# SSH 登录服务器\n').italic = True
    code1.add_run('ssh root@47.103.108.47\n\n')
    code1.add_run('# 启动服务\n').italic = True
    code1.add_run('systemctl start gunicorn-flask-data-server\n\n')
    code1.add_run('# 查看服务状态\n').italic = True
    code1.add_run('systemctl status gunicorn-flask-data-server\n\n')
    code1.add_run('# 设置开机自启\n').italic = True
    code1.add_run('systemctl enable gunicorn-flask-data-server')
    
    doc.add_paragraph()
    doc.add_paragraph('方法 2：直接使用 Gunicorn 启动', style='Heading 3')
    code2 = doc.add_paragraph()
    code2.add_run('cd /root/course-project/week5/data-server\n')
    code2.add_run('/usr/local/bin/gunicorn -c config/gunicorn_config.py app:app -D')
    
    doc.add_paragraph()
    doc.add_paragraph('方法 3：使用 Flask 开发服务器（仅测试）', style='Heading 3')
    code3 = doc.add_paragraph()
    code3.add_run('cd /root/course-project/week5/data-server\n')
    code3.add_run('python3 app.py')
    
    doc.add_heading('1.3 验证服务启动成功', level=2)
    
    verify_steps = [
        ('步骤 1', '检查进程'),
        ('步骤 2', '检查端口监听'),
        ('步骤 3', '测试健康检查接口'),
        ('步骤 4', '测试数据接收接口')
    ]
    
    for step, desc in verify_steps:
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{step}：{desc}').bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('验证命令：', style='Heading 3')
    verify_code = doc.add_paragraph()
    verify_code.add_run('# 1. 检查 Gunicorn 进程\n')
    verify_code.add_run('ps aux | grep gunicorn\n\n')
    verify_code.add_run('# 2. 检查端口监听\n')
    verify_code.add_run('netstat -tlnp | grep 5000\n\n')
    verify_code.add_run('# 3. 测试健康检查\n')
    verify_code.add_run('curl http://localhost:5000/api/health\n\n')
    verify_code.add_run('# 4. 测试数据接收\n')
    verify_code.add_run('curl -X POST http://localhost:5000/api/receive \\\n')
    verify_code.add_run('  -H "Content-Type: application/json" \\\n')
    verify_code.add_run('  -d \'{"sensor_type":"skin","moisture":65,"oiliness":32}\'')
    
    doc.add_page_break()
    
    # 第二部分：鉴权机制
    doc.add_heading('二、鉴权机制详解', level=1)
    
    doc.add_heading('2.1 什么是鉴权？', level=2)
    p = doc.add_paragraph()
    p.add_run('鉴权（Authentication & Authorization）').bold = True
    p.add_run('是验证用户身份并授权访问的过程。就像去公司需要刷卡进门一样，API 也需要验证请求者的身份。')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('为什么需要鉴权？').bold = True
    reasons = [
        '防止未授权用户访问敏感数据',
        '保护服务器资源不被滥用',
        '记录操作日志，追踪用户行为',
        '实现不同角色的权限控制'
    ]
    for reason in reasons:
        doc.add_paragraph(reason, style='List Bullet')
    
    doc.add_heading('2.2 你的系统支持的鉴权方式', level=2)
    
    doc.add_heading('方式 1：JWT Token 认证（推荐）', level=3)
    
    p = doc.add_paragraph()
    p.add_run('JWT（JSON Web Token）').bold = True
    p.add_run('是一种开放标准，用于在各方之间安全地传输信息。')
    
    doc.add_paragraph()
    doc.add_paragraph('工作原理：', style='Heading 4')
    jwt_steps = [
        '客户端发送用户名和密码到 /api/auth/login',
        '服务器验证凭据，生成 JWT Token',
        '客户端在后续请求中携带 Token（Authorization: Bearer <token>）',
        '服务器验证 Token 有效性，允许或拒绝访问'
    ]
    for i, step in enumerate(jwt_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph('演示命令：', style='Heading 4')
    jwt_demo = doc.add_paragraph()
    jwt_demo.add_run('# 步骤 1：登录获取 Token\n')
    jwt_demo.add_run('curl -X POST http://localhost:5000/api/auth/login \\\n')
    jwt_demo.add_run('  -H "Content-Type: application/json" \\\n')
    jwt_demo.add_run('  -d \'{"username":"admin","password":"admin123"}\'\n\n')
    jwt_demo.add_run('# 返回：\n')
    jwt_demo.add_run('# {\n')
    jwt_demo.add_run('#   "token": "eyJ0eXAiOiJKV1QiLCJhbGci...",\n')
    jwt_demo.add_run('#   "status": "success"\n')
    jwt_demo.add_run('# }\n\n')
    jwt_demo.add_run('# 步骤 2：使用 Token 访问受保护接口\n')
    jwt_demo.add_run('curl -X POST http://localhost:5000/api/receive/secure \\\n')
    jwt_demo.add_run('  -H "Content-Type: application/json" \\\n')
    jwt_demo.add_run('  -H "Authorization: Bearer eyJ0eXAiOiJKV1QiLCJhbGci..." \\\n')
    jwt_demo.add_run('  -d \'{"encrypted_data":"..."}\'')
    
    doc.add_paragraph()
    doc.add_heading('方式 2：API Key 认证', level=3)
    
    apikey_demo = doc.add_paragraph()
    apikey_demo.add_run('# 使用 API Key 访问\n')
    apikey_demo.add_run('curl -X POST http://localhost:5000/api/receive/apikey \\\n')
    apikey_demo.add_run('  -H "Content-Type: application/json" \\\n')
    apikey_demo.add_run('  -H "X-API-Key: your_api_key_here" \\\n')
    apikey_demo.add_run('  -d \'{"sensor_type":"skin","moisture":65}\'')
    
    doc.add_heading('2.3 鉴权验证演示（给老师看）', level=2)
    
    doc.add_paragraph('演示 1：无鉴权访问被拒绝', style='Heading 3')
    demo1 = doc.add_paragraph()
    demo1.add_run('# 尝试不使用 Token 访问受保护接口\n')
    demo1.add_run('curl -X POST http://localhost:5000/api/receive/secure \\\n')
    demo1.add_run('  -H "Content-Type: application/json" \\\n')
    demo1.add_run('  -d \'{"encrypted_data":"test"}\'\n\n')
    demo1.add_run('# 预期返回：401 Unauthorized\n')
    demo1.add_run('# {\n')
    demo1.add_run('#   "error": "未提供认证信息"\n')
    demo1.add_run('# }')
    
    doc.add_paragraph()
    doc.add_paragraph('演示 2：正确鉴权访问成功', style='Heading 3')
    demo2 = doc.add_paragraph()
    demo2.add_run('# 1. 先登录获取 Token\n')
    demo2.add_run('TOKEN=$(curl -s -X POST http://localhost:5000/api/auth/login \\\n')
    demo2.add_run('  -H "Content-Type: application/json" \\\n')
    demo2.add_run('  -d \'{"username":"admin","password":"admin123"}\' \\\n')
    demo2.add_run('  | grep -o \'"token":"[^"]*\'"\' | cut -d\'"\' -f4)\n\n')
    demo2.add_run('echo "获取到 Token: $TOKEN"\n\n')
    demo2.add_run('# 2. 使用 Token 访问（应该成功）\n')
    demo2.add_run('curl -X POST http://localhost:5000/api/receive/secure \\\n')
    demo2.add_run('  -H "Content-Type: application/json" \\\n')
    demo2.add_run('  -H "Authorization: Bearer $TOKEN" \\\n')
    demo2.add_run('  -d \'{"encrypted_data":"test"}\'\n\n')
    demo2.add_run('# 预期返回：200 OK')
    
    doc.add_paragraph()
    doc.add_paragraph('演示 3：无效 Token 被拒绝', style='Heading 3')
    demo3 = doc.add_paragraph()
    demo3.add_run('# 使用无效的 Token\n')
    demo3.add_run('curl -X POST http://localhost:5000/api/receive/secure \\\n')
    demo3.add_run('  -H "Content-Type: application/json" \\\n')
    demo3.add_run('  -H "Authorization: Bearer invalid_token_here" \\\n')
    demo3.add_run('  -d \'{"encrypted_data":"test"}\'\n\n')
    demo3.add_run('# 预期返回：401 Unauthorized\n')
    demo3.add_run('# {\n')
    demo3.add_run('#   "error": "无效的 Token"\n')
    demo3.add_run('# }')
    
    doc.add_page_break()
    
    # 第三部分：协议文档
    doc.add_heading('三、协议文档说明', level=1)
    
    doc.add_heading('3.1 什么是协议文档？', level=2)
    p = doc.add_paragraph()
    p.add_run('协议文档（API Documentation）').bold = True
    p.add_run('是 API 的使用说明书，告诉使用者：')
    protocols = [
        '接口地址是什么',
        '需要什么参数（请求格式）',
        '返回什么数据（响应格式）',
        '如何调用（示例代码）',
        '错误码说明'
    ]
    for proto in protocols:
        doc.add_paragraph(proto, style='List Bullet')
    
    doc.add_heading('3.2 你的系统提供的协议文档', level=2)
    
    doc_table = doc.add_table(rows=5, cols=3)
    doc_table.style = 'Table Grid'
    doc_headers = ['文档类型', '文件位置', '用途']
    for i, header in enumerate(doc_headers):
        doc_table.rows[0].cells[i].text = header
        doc_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    doc_data = [
        ['Markdown', 'API 接口文档（）.md', '主要接口文档，包含所有接口说明'],
        ['Markdown', 'API 文档.md', '详细说明文档'],
        ['Swagger/OpenAPI', 'swagger.json', '可导入 Swagger UI 在线测试'],
        ['Postman', 'postman_collection.json', '可导入 Postman 可视化测试']
    ]
    
    for i, row_data in enumerate(doc_data, 1):
        for j, cell_data in enumerate(row_data):
            doc_table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_heading('3.3 如何查看和使用协议文档', level=2)
    
    doc.add_paragraph('方法 1：直接查看 Markdown 文档', style='Heading 3')
    method1 = doc.add_paragraph()
    method1.add_run('# 在服务器上\n')
    method1.add_run('cat /root/course-project/week5/data-server/API\\ 接口文档（）.md\n\n')
    method1.add_run('# 或在本地用 VSCode 打开\n')
    method1.add_run('d:\\学习\\软件设计\\data-server\\docs\\API 接口文档（）.md')
    
    doc.add_paragraph()
    doc.add_paragraph('方法 2：使用 Postman 可视化测试', style='Heading 3')
    method2 = doc.add_paragraph()
    method2.add_run('1. 打开 Postman\n')
    method2.add_run('2. 点击 Import → 选择 postman_collection.json\n')
    method2.add_run('3. 导入后即可看到所有接口的可视化界面\n')
    method2.add_run('4. 点击 Send 按钮即可测试接口')
    
    doc.add_paragraph()
    doc.add_paragraph('方法 3：使用 Swagger UI', style='Heading 3')
    method3 = doc.add_paragraph()
    method3.add_run('1. 访问 https://swagger.io/tools/swagger-ui/\n')
    method3.add_run('2. 点击 "Explore" → 输入 swagger.json 的 URL\n')
    method3.add_run('3. 可以在线测试所有接口')
    
    doc.add_heading('3.4 协议文档内容示例', level=2)
    
    example_table = doc.add_table(rows=8, cols=2)
    example_table.style = 'Table Grid'
    
    example_headers = ['字段', '说明']
    for i, header in enumerate(example_headers):
        example_table.rows[0].cells[i].text = header
        example_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    example_data = [
        ['接口名称', '数据接收接口'],
        ['接口地址', 'POST /api/receive'],
        ['请求格式', 'Content-Type: application/json'],
        ['请求参数', '{"sensor_type":"skin","moisture":65,"oiliness":32}'],
        ['响应格式', '{"status":"success","message":"数据接收成功","filename":"data_xxx.json"}'],
        ['认证方式', '无需认证（公开接口）'],
        ['错误码', '400: 参数错误, 500: 服务器错误']
    ]
    
    for i, row_data in enumerate(example_data, 1):
        for j, cell_data in enumerate(row_data):
            example_table.rows[i].cells[j].text = cell_data
    
    doc.add_page_break()
    
    # 第四部分：本地模拟器演示
    doc.add_heading('四、本地模拟器演示流程', level=1)
    
    doc.add_heading('4.1 演示目标', level=2)
    p = doc.add_paragraph()
    p.add_run('向老师展示：').bold = True
    p.add_run('本地模拟器产生数据 → 通过 HTTP API 发送 → 服务器接收并保存')
    
    doc.add_heading('4.2 演示前准备', level=2)
    
    prep_items = [
        '确保服务器 Gunicorn 服务已启动',
        '确保本地可以 SSH 连接到服务器',
        '准备好模拟器脚本（simulator_mq.py）',
        '打开终端，分两个窗口（一个本地，一个服务器）'
    ]
    for item in prep_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.3 完整演示步骤', level=2)
    
    doc.add_heading('步骤 1：检查服务器状态', level=3)
    step1 = doc.add_paragraph()
    step1.add_run('# 在服务器终端执行\n')
    step1.add_run('ssh root@47.103.108.47\n\n')
    step1.add_run('# 检查服务是否运行\n')
    step1.add_run('ps aux | grep gunicorn\n\n')
    step1.add_run('# 如果服务未启动，执行：\n')
    step1.add_run('systemctl start gunicorn-flask-data-server\n\n')
    step1.add_run('# 测试健康检查\n')
    step1.add_run('curl http://localhost:5000/api/health')
    
    doc.add_paragraph()
    doc.add_paragraph('预期输出：', style='Heading 4')
    output1 = doc.add_paragraph()
    output1.add_run('{\n')
    output1.add_run('  "status": "healthy",\n')
    output1.add_run('  "service": "Flask Data Server v3.0"\n')
    output1.add_run('}')
    
    doc.add_paragraph()
    doc.add_heading('步骤 2：启动本地模拟器', level=3)
    step2 = doc.add_paragraph()
    step2.add_run('# 在本地终端执行\n')
    step2.add_run('cd d:\\学习\\软件设计\\data-server\n\n')
    step2.add_run('# 运行模拟器（会产生皮肤和环境数据）\n')
    step2.add_run('python3 simulator_mq.py\n\n')
    step2.add_run('# 预期输出：\n')
    step2.add_run('# 🚀 魔镜数据模拟器 v2.0\n')
    step2.add_run('# [11:36:23] 💧 皮肤数据产生 -> 水分:68% | 油亮:36%\n')
    step2.add_run('#    💧 [发布成功] MQ接收: skin (11:36:23)')
    
    doc.add_paragraph()
    doc.add_heading('步骤 3：在服务器端查看数据接收情况', level=3)
    step3 = doc.add_paragraph()
    step3.add_run('# 在服务器终端执行\n')
    step3.add_run('# 查看最新的数据文件\n')
    step3.add_run('ls -lt /root/data-server/data/ | head -10\n\n')
    step3.add_run('# 查看数据文件总数\n')
    step3.add_run('find /root/data-server/data/ -name "*.json" | wc -l\n\n')
    step3.add_run('# 查看最新文件内容\n')
    step3.add_run('ls -t /root/data-server/data/*.json | head -1 | xargs cat | python3 -m json.tool')
    
    doc.add_paragraph()
    doc.add_paragraph('预期输出：', style='Heading 4')
    output2 = doc.add_paragraph()
    output2.add_run('{\n')
    output2.add_run('  "sensor_type": "skin",\n')
    output2.add_run('  "moisture": 68,\n')
    output2.add_run('  "oiliness": 36,\n')
    output2.add_run('  "timestamp": "2026-04-17T11:36:23"\n')
    output2.add_run('}')
    
    doc.add_paragraph()
    doc.add_heading('步骤 4：手动测试单个数据上传', level=3)
    step4 = doc.add_paragraph()
    step4.add_run('# 在本地或服务器终端执行\n')
    step4.add_run('curl -X POST http://47.103.108.47:5000/api/receive \\\n')
    step4.add_run('  -H "Content-Type: application/json" \\\n')
    step4.add_run('  -d \'{"sensor_type":"skin","moisture":65,"oiliness":32,"timestamp":"2026-04-17 12:00:00"}\'\n\n')
    step4.add_run('# 预期返回：\n')
    step4.add_run('# {\n')
    step4.add_run('#   "status": "success",\n')
    step4.add_run('#   "message": "数据接收成功",\n')
    step4.add_run('#   "filename": "data_20260417_120000_xxx.json"\n')
    step4.add_run('# }')
    
    doc.add_paragraph()
    doc.add_heading('步骤 5：展示数据已保存', level=3)
    step5 = doc.add_paragraph()
    step5.add_run('# 在服务器终端执行\n')
    step5.add_run('# 查看刚才上传的文件\n')
    step5.add_run('cat /root/data-server/data/data_20260417_120000_xxx.json | python3 -m json.tool')
    
    doc.add_heading('4.4 演示话术建议', level=2)
    
    speech = doc.add_paragraph()
    speech.add_run('"现在我来演示本地模拟器如何将数据发送到服务器：\n\n')
    speech.add_run('1. 首先，我在本地运行模拟器，它会产生模拟的皮肤传感器数据。\n')
    speech.add_run('2. 模拟器通过 HTTP POST 请求，调用服务器的 /api/receive 接口。\n')
    speech.add_run('3. 服务器接收到数据后，将其保存为 JSON 文件。\n')
    speech.add_run('4. 大家可以看到，数据已经成功保存到服务器的 /root/data-server/data/ 目录下。\n')
    speech.add_run('5. 这就是我们系统的核心功能：模拟数据 → HTTP 传输 → 服务器接收 → 本地保存。"')
    
    doc.add_page_break()
    
    # 第五部分：完整验收演示脚本
    doc.add_heading('五、完整验收演示脚本', level=1)
    
    doc.add_heading('5.1 快速启动服务', level=2)
    quick_start = doc.add_paragraph()
    quick_start.add_run('# 一键启动服务脚本\n')
    quick_start.add_run('ssh root@47.103.108.47 << \'EOF\'\n')
    quick_start.add_run('# 启动 Gunicorn 服务\n')
    quick_start.add_run('systemctl start gunicorn-flask-data-server\n')
    quick_start.add_run('sleep 3\n\n')
    quick_start.add_run('# 验证服务启动\n')
    quick_start.add_run('echo "=== 服务状态 ==="\n')
    quick_start.add_run('systemctl is-active gunicorn-flask-data-server\n\n')
    quick_start.add_run('echo "=== 健康检查 ==="\n')
    quick_start.add_run('curl -s http://localhost:5000/api/health | python3 -m json.tool\n\n')
    quick_start.add_run('echo "=== 数据文件数量 ==="\n')
    quick_start.add_run('find /root/data-server/data/ -name "*.json" | wc -l\n')
    quick_start.add_run('EOF')
    
    doc.add_heading('5.2 完整验收流程', level=2)
    full_demo = doc.add_paragraph()
    full_demo.add_run('# 完整的验收演示流程（约 10 分钟）\n\n')
    full_demo.add_run('ssh root@47.103.108.47 << \'DEMO\'\n\n')
    full_demo.add_run('echo "========================================"\n')
    full_demo.add_run('echo "  第二阶段验收演示"\n')
    full_demo.add_run('echo "========================================"\n\n')
    
    demo_steps = [
        ('1', '系统架构', '技术栈、服务状态'),
        ('2', 'RESTful API', '健康检查、接口列表'),
        ('3', '数据上传', '模拟器演示、手动上传'),
        ('4', '接口文档', '展示文档文件'),
        ('5', '鉴权功能', 'Token 获取、拒绝、通过'),
        ('6', '数据加密', '加密、传输、解密'),
        ('7', '消息队列', 'Redis、模拟器 MQ 模式'),
        ('8', '性能指标', 'QPS、成功率')
    ]
    
    for num, title, desc in demo_steps:
        full_demo.add_run(f'\necho "【{num}】{title} - {desc}"\n')
        full_demo.add_run(f'echo "----------------------------------------"\n')
        full_demo.add_run(f'# 具体命令...\n\n')
    
    full_demo.add_run('echo "========================================"\n')
    full_demo.add_run('echo "  验收完成！"\n')
    full_demo.add_run('echo "========================================"\n')
    full_demo.add_run('DEMO')
    
    doc.add_page_break()
    
    # 第六部分：常见问题
    doc.add_heading('六、常见问题与解决方案', level=1)
    
    faq_table = doc.add_table(rows=8, cols=2)
    faq_table.style = 'Table Grid'
    
    faq_headers = ['问题', '解决方案']
    for i, header in enumerate(faq_headers):
        faq_table.rows[0].cells[i].text = header
        faq_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    faq_data = [
        ('Connection refused', '服务未启动，执行 systemctl start gunicorn-flask-data-server'),
        ('401 Unauthorized', '未提供 Token，先调用 /api/auth/login 获取'),
        ('403 Forbidden', 'Token 无效或过期，重新登录获取新 Token'),
        ('400 Bad Request', '请求参数格式错误，检查 Content-Type 和 JSON 格式'),
        ('429 Too Many Requests', '触发限流，等待 1 分钟后重试'),
        ('500 Internal Error', '服务器内部错误，查看日志 tail -f logs/gunicorn_error.log'),
        ('Redis 连接失败', '检查 Redis 服务 systemctl status redis')
    ]
    
    for i, (question, solution) in enumerate(faq_data, 1):
        faq_table.rows[i].cells[0].text = question
        faq_table.rows[i].cells[1].text = solution
    
    # 保存文档
    output_path = 'd:\\学习\\软件设计\\data-server\\docs\\第二阶段完整验收指南.docx'
    doc.save(output_path)
    print(f'✅ Word 文档已生成：{output_path}')

if __name__ == '__main__':
    create_complete_guide()
