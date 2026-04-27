#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成第二阶段验收指南 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_acceptance_guide():
    # 创建文档
    doc = Document()
    
    # 标题
    title = doc.add_heading('Flask 数据服务器第二阶段验收指南', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('一、项目基本信息', level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ['服务器地址', '47.103.108.47:5000'],
        ['技术栈', 'Flask + Gunicorn + Redis + gevent'],
        ['Python 版本', '3.6.8'],
        ['Worker 配置', '7 workers (gevent)'],
        ['验收完成度', '7/8 项（HTTPS 为可选）']
    ]
    
    for i, (key, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = key
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    
    # 验收内容
    doc.add_heading('二、验收内容详解', level=1)
    
    # (1) RESTful API
    doc.add_heading('（1）RESTful API 接口', level=2)
    doc.add_paragraph('状态：已完成')
    doc.add_paragraph('提供 12 个基于 REST 风格的 API 接口：')
    
    api_table = doc.add_table(rows=13, cols=3)
    api_table.style = 'Table Grid'
    api_table.rows[0].cells[0].text = 'HTTP 方法'
    api_table.rows[0].cells[1].text = '接口路径'
    api_table.rows[0].cells[2].text = '功能'
    
    for cell in api_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    apis = [
        ['GET', '/api/health', '健康检查'],
        ['POST', '/api/receive', '数据接收'],
        ['POST', '/api/sensor/skin', '皮肤数据'],
        ['POST', '/api/sensor/environment', '环境数据'],
        ['POST', '/api/device/status', '设备状态'],
        ['POST', '/api/auth/login', '登录认证'],
        ['POST', '/api/auth/apikey', 'API Key'],
        ['POST', '/api/receive/secure', '加密上传'],
        ['POST', '/api/encrypt', '数据加密'],
        ['POST', '/api/decrypt', '数据解密'],
        ['POST', '/api/receive/apikey', 'API Key 认证'],
        ['GET', '/api/stats', '统计信息']
    ]
    
    for i, (method, path, func) in enumerate(apis, 1):
        api_table.rows[i].cells[0].text = method
        api_table.rows[i].cells[1].text = path
        api_table.rows[i].cells[2].text = func
    
    doc.add_paragraph()
    doc.add_paragraph('演示命令：')
    doc.add_paragraph('curl http://localhost:5000/api/health')
    
    # (2) 数据上传
    doc.add_heading('（2）模拟数据上传与本地保存', level=2)
    doc.add_paragraph('状态：已完成')
    doc.add_paragraph('实现方式：')
    doc.add_paragraph('1. 模拟器（simulator_mq.py）产生传感器数据')
    doc.add_paragraph('2. 通过 REST API 发送到服务器')
    doc.add_paragraph('3. 服务器接收数据并保存到本地文件系统')
    doc.add_paragraph('4. 数据格式：JSON，按时间戳命名')
    
    doc.add_paragraph()
    doc.add_paragraph('验证方法：')
    doc.add_paragraph('# 查看数据文件数量')
    doc.add_paragraph('find /root/data-server/data/ -name "*.json" | wc -l')
    doc.add_paragraph('# 输出：443506')
    
    # (3) HTTPS
    doc.add_heading('（3）HTTPS 访问支持', level=2)
    doc.add_paragraph('状态：可选功能（未配置）')
    doc.add_paragraph('说明：由于未购买 SSL 证书，HTTPS 功能暂未实现。但系统架构已预留 SSL 配置接口，后续可通过配置 Nginx 反向代理快速启用。')
    
    # (4) 接口文档
    doc.add_heading('（4）接口文档', level=2)
    doc.add_paragraph('状态：已完成')
    doc.add_paragraph('文档位置：')
    doc.add_paragraph('/root/data-server/API 接口文档（）.md - Markdown 格式')
    doc.add_paragraph('/root/data-server/swagger.json - OpenAPI 规范')
    doc.add_paragraph('/root/data-server/postman_collection.json - Postman 集合')
    
    # (5) 鉴权
    doc.add_heading('（5）鉴权功能', level=2)
    doc.add_paragraph('状态：已完成')
    doc.add_paragraph('支持两种鉴权方式：')
    
    doc.add_paragraph('方式 1：JWT Token 认证')
    doc.add_paragraph('# 登录获取 Token')
    doc.add_paragraph('curl -X POST http://localhost:5000/api/auth/login \\')
    doc.add_paragraph('  -d \'{\"username\":\"admin\",\"password\":\"admin123\"}\'')
    
    doc.add_paragraph()
    doc.add_paragraph('方式 2：API Key 认证')
    doc.add_paragraph('curl -X POST http://localhost:5000/api/receive/apikey \\')
    doc.add_paragraph('  -H "X-API-Key: your_api_key_here"')
    
    doc.add_paragraph()
    doc.add_paragraph('鉴权验证：')
    doc.add_paragraph('# 无 Token 访问（应返回 401）')
    doc.add_paragraph('curl -X POST http://localhost:5000/api/receive/secure -d \'{}\'')
    
    # (6) 数据加密
    doc.add_heading('（6）数据加密传输', level=2)
    doc.add_paragraph('状态：已完成')
    doc.add_paragraph('加密算法：AES-256（Fernet 对称加密）')
    doc.add_paragraph('实现流程：')
    doc.add_paragraph('1. 客户端调用 /api/encrypt 加密数据')
    doc.add_paragraph('2. 传输加密后的数据到 /api/receive/secure')
    doc.add_paragraph('3. 服务器自动解密并保存到本地')
    
    # (7) 消息队列
    doc.add_heading('（7）消息队列', level=2)
    doc.add_paragraph('状态：已完成')
    doc.add_paragraph('消息队列：Redis Streams')
    doc.add_paragraph('实现方式：')
    doc.add_paragraph('1. 模拟器通过 Redis 发布消息到 sensor:raw 队列')
    doc.add_paragraph('2. 消费者监听队列并处理数据')
    doc.add_paragraph('3. 处理后的数据保存到本地文件系统')
    
    doc.add_paragraph()
    doc.add_paragraph('演示命令：')
    doc.add_paragraph('redis-cli ping  # 返回 PONG')
    doc.add_paragraph('redis-cli keys \'sensor:*\'')
    doc.add_paragraph('python3 simulator_mq.py')
    
    # 总结
    doc.add_heading('三、验收总结', level=1)
    
    summary_table = doc.add_table(rows=8, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_headers = ['验收项', '状态']
    for i, header in enumerate(summary_headers):
        summary_table.rows[0].cells[i].text = header
        summary_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    summary_data = [
        ['(1) RESTful API', '已完成'],
        ['(2) 数据上传保存', '已完成'],
        ['(3) HTTPS', '可选（未配置）'],
        ['(4) 接口文档', '已完成'],
        ['(5) 鉴权功能', '已完成'],
        ['(6) 数据加密', '已完成'],
        ['(7) 消息队列', '已完成']
    ]
    
    for i, row_data in enumerate(summary_data, 1):
        for j, cell_data in enumerate(row_data):
            summary_table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    doc.add_paragraph('总体完成度：7/8 项（核心功能 100%）')
    
    # 保存
    output_path = 'd:\\学习\\软件设计\\data-server\\docs\\第二阶段验收指南.docx'
    doc.save(output_path)
    print(f'✅ Word 文档已生成：{output_path}')

if __name__ == '__main__':
    create_acceptance_guide()
